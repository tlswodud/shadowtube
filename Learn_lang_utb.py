#!/usr/bin/env python
# coding: utf-8

# In[1]:


import torch

print("Number of GPU: ", torch.cuda.device_count())
print("GPU Name: ", torch.cuda.get_device_name())


device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
print('Using device:', device)


# In[2]:


from urllib.parse import urlparse, parse_qs

def get_video_id(url):
    parsed_url = urlparse(url)
    
    # 긴 형식 (https://www.youtube.com/watch?v=영상아이디)인 경우
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path == '/watch':
        return parse_qs(parsed_url.query).get('v', [None])[0]
    
    # 짧은 형식 (https://youtu.be/영상아이디)인 경우
    if parsed_url.hostname == 'youtu.be':
        return parsed_url.path[1:]  # 첫 번째 슬래시 제거
    
    # 임베드 형식 (https://www.youtube.com/embed/영상아이디)인 경우
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path.startswith('/embed/'):
        return parsed_url.path.split('/')[2]
    
    return None

# 예시 URL
#url = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"

if __name__ == "__main__":
    url = input()
    video_id = get_video_id(url)
    print(video_id)  # 출력: dQw4w9WgXcQ


# In[3]:


from youtube_transcript_api import YouTubeTranscriptApi

transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)


# In[4]:


def list_available_languages(video_id):
    # 특정 비디오의 자막 언어 목록을 가져옵니다.
    #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
    
    # 사용 가능한 언어 코드와 이름을 출력합니다.
    languages = {}
    for transcript_item in transcript_list:
        languages[transcript_item.language_code] = transcript_item.language
    
    return languages


def check_languages(video_id): # 영어 목록들 가져오는것 # 영어 자막이 없다면 중지
    
    preferred_languages = ['en-US', 'en-GB', 'en-CA', 'en-AU', 'en' ,'a.en']
    
    list_lang = list_available_languages(video_id)
    
    # 언어 코드 리스트에 대해 확인
    for lang_code in list_lang:
        if lang_code in preferred_languages:
            print(f"해당 언어 자막 확인 완료했습니다 잠시만 기다려주세요: {list_lang[lang_code]}")
            return lang_code

    for lang_code in list_lang:
        if lang_code.startswith("en"):  # 'en'으로 시작하는 모든 코드 체크
            print(f"해당 언어 자막 확인 완료했습니다 잠시만 기다려주세요: {list_lang[lang_code]}")
            return lang_code   

    # 원하는 언어가 없을 경우
    print("해당 자막이 존재하지 않습니다 다른 영상을 선택해주세요")
    return None 

if __name__ == "__main__":
    list_lang = list_available_languages(video_id)

    for i in list_lang:
        print(i)  

    en_coder = check_languages(video_id)



# In[5]:


def ko_isavailable(video_id):
    #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

    for transcript in transcript_list:
        transcript.is_translatable
        print("한국어 번역 자막 사용가능합니다")
        return transcript.is_translatable

def get_fetch(video_id):

    #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
    for transcript in transcript_list:
        transcript.is_translated
        
if __name__ == "__main__":
    ko_isavailable(video_id)


# In[6]:


def check_languages_ko(video_id):
    # 사용할 언어 목록을 미리 정의합니다.
    preferred_languages = ['a.ko']
    
    list_lang = list_available_languages(video_id)
    
    # 언어 코드 리스트에 대해 확인
    for lang_code in list_lang:
        if lang_code in preferred_languages:
            print(f"만약 번역이 없을경우 자동완성은 있네요: {list_lang[lang_code]}")
            return lang_code 


    # 원하는 언어가 없을 경우
    print("해당 자막이 존재하지 않습니다 다른 영상을 선택해주세요")
    return None 
if __name__ == "__main__":

    list_lang = list_available_languages(video_id)

    for i in list_lang:
        print(i)  


    if ko_isavailable(video_id) == False:
        
        ko_coder = check_languages_ko(video_id)
        print(ko_coder)


# In[7]:


from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document

def get_transcript(video_id):
    # YouTube Transcript API로 자막 가져오기
    transcript_from = YouTubeTranscriptApi.get_transcript(video_id, languages=[en_coder])  # 영어와 한국어 자막 가져옴
    # 자막 텍스트만 추출하여 리스트로 만듭니다.
    transcript_texts = [(entry['start'], entry['text'])  for entry in transcript_from]  # 시작 및 종료 시간 정보를 제외하고 텍스트만 추출
    return transcript_texts  # 텍스트 리스트 반환

# def get_transcript_ko(video_id):
#     # YouTube Transcript API로 자막 가져오고 번역본을 들고오는 것
#     #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
    
#     for transcript in transcript_list:
#         translated_transcript = transcript.translate('ko')# 한국어로 변환

#     transcript_texts = [entry['text'] for entry in translated_transcript.fetch()]  # 시작 및 종료 시간 정보를 제외하고 텍스트만 추출
#     return transcript_texts  # 텍스트 리스트 반환

def contains_lowercase(video_id):
    
    transcript_from  = YouTubeTranscriptApi.get_transcript(video_id, languages=[en_coder])
    transcript_texts = [entry['text'] for entry in transcript_from]  # 시작 및 종료 시간 정보를 제외하고 텍스트만 추출
    
    return any(c.islower() for c in transcript_texts)  

# 유튜브 영상 ID
def check_dot(video_id):
    transcript_from  = YouTubeTranscriptApi.get_transcript(video_id, languages=[en_coder])
    transcript_texts = [entry['text'] for entry in transcript_from]

    for test_dot in transcript_texts:

        if '.' in test_dot:
        
            return True
    return False    


# In[8]:

if __name__ == "__main__":
    en_coder


# In[9]:
if __name__ == "__main__":

    LowerOrUpper =  contains_lowercase(video_id)        

    if LowerOrUpper == True:
        print("소문자가 있다")

    else:
        print("소문자가 없다")   

    if check_dot(video_id) == True:
        print(". 이 있습니다")

    else:
        print(". 이 없습니다")   


# In[10]:


import yt_dlp

def get_video_title(video_url):
    # yt-dlp 객체 생성
    ydl_opts = {}
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        # 비디오 정보 추출
        info_dict = ydl.extract_info(video_url, download=False)
        # 제목 가져오기
        title = info_dict.get('title', None)
        return title
if __name__ == "__main__":    
    utb_title = get_video_title(video_id)
    utb_title


# **유튜브 자동완성을 통한 번역본을 가져왔음** 
# 
# . 이 예상과 다르게 찍힌 부분이 있어 번갈아 가며 출력이 어렵다.
# 
# 따로 제공하되 AI 를 통한 단어 제공 정도는 가능성 있다 
# 
# 정 번역 기능 추가가 어렵다면 말이다.
# 
# **사용할수 있지만 완벽하지는 않은 듯 하다**
# 
# 속도? 당연히 빠르다 
# 
# 해석? 안되는 부분도 없다 
# 
# 다만  translate 기능에 비해 아쉽다는거다
# 
# openai 또는 제미니 Api를 테스트 해야할듯
# 

# In[11]:

if __name__ == "__main__":
    transcript_list = get_transcript(video_id)
    #translator_target = get_transcript_ko(video_id)


# In[12]:


import re 

new_script = ""

for start, read_script in transcript_list:
        minutes = int(start // 60)  # 분 계산 (소수점 없음)
        seconds = int(start % 60)  # 초 계산 (소수점 없음)
    
    # 시간 형식 설정 (분.초 형태)
        time_format = f"[{minutes:02d}:{seconds:02d}]"
    # . 기반이다 보니 문제가 있을 만한 것들을 수정    
        read_script = read_script.replace('U.S.', 'US')
        read_script = read_script.replace('U.S', 'US')
        
        read_script = read_script.replace('Mr.', 'Mr ')
        read_script = read_script.replace('Mrs.', 'Mrs ')

        read_script = read_script.replace('Ph.D.', 'ph,D ')
        read_script = read_script.replace('Prof.', 'prof ')
        read_script = read_script.replace('Dr.', 'Dr ')

        read_script = read_script.replace('No.', 'No,')
        
        read_script = read_script.replace('a.m.', 'am')
        read_script = read_script.replace('p.m.', 'pm')
        
        read_script = re.sub(r'(\d)\.(\d)', r'\1_\2', read_script)


        read_script = read_script.replace('\n', ' ')
        read_script = read_script.replace('.', '. \n')
        read_script = read_script.replace('?' , '? \n')
        read_script = read_script.replace('>>' ,'\n >>')
        if LowerOrUpper == False:
                read_script = read_script[0].upper() + read_script[1:].lower()

        new_script += ' '
        new_script += time_format
        new_script += read_script



# In[13]:


import re
result_transcript =  ["\n\n"]

to_timestamps_list = []

def clean_transcript_texts(transcript_texts):
    cleaned_texts = ""
    for text in transcript_texts:
        # 첫 번째 타임스탬프만 남기고 나머지 타임스탬프를 제거
        # 1) 모든 타임스탬프를 찾음
        timestamps = re.findall(r'\[\d{2}:\d{2}\]', text)
        
        if timestamps:
            # 2) 첫 번째 타임스탬프만 남기고  리스트에 to_time  에 넣어주었음 나중에 앞에 붙일거
            first_timestamp = timestamps[0]
            to_timestamps_list.append(first_timestamp)
            cleaned_text = text.replace(first_timestamp, '',1)
            cleaned_text = re.sub(r'\[\d{2}:\d{2}\]','', cleaned_text)  # 나머지 타임스탬프 제거
            cleaned_texts += first_timestamp +" "+ cleaned_text.strip() +" "
           
        else:
            # 타임스탬프가 없는 경우
            cleaned_texts += text.strip() +" "

    return cleaned_texts.strip()  # 최종 문자열 반환

new_script_line = new_script.splitlines()

for line in new_script_line:
    result_transcript.append("\n")
    result_transcript.append(clean_transcript_texts([line]))
    result_transcript.append("\n")



# In[14]:

if __name__ == "__main__":
    doc = Document()
    doc.add_heading(f'{utb_title} YouTube Transcript', level=1)  # 문서 제목 추가
    doc.add_paragraph(result_transcript)  # 스크립트 추가
    # 문서 저장
    doc.save('transcript_v6_only_eng.docx')
    print("transcript_v6_only_eng.docx로 저장되었습니다.")


# **앞에 start 넣어주고 이걸 기반으로 ko 찾아보자 그럼 .을 기반으로 한 번역은 필요가 없어진다**
# 사실 번역기 쓰는게 제일 편할텐데 ㅋㅋㅋ
# 
# * 실패 불가능함 인덱스라 생각한 시작 지점이 엉망이다
# 

# In[15]:
if __name__ == "__main__":

    result_transcript 


# **여기까지가 영어 [숫자] 표기 해준거임**
# 
# 자동완성으로 만든 한글말고 번역을 써주는 코드 

# In[16]:


transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

translator_target=""

for transcript in transcript_list:
    
    translator_target += " ".join([item['text'] for item in transcript.translate('ko').fetch()]) + "\n\n"


# In[17]:


translator_target = translator_target.splitlines()


# In[18]:


new_script_target = ""

for read_script_target in translator_target:
        
        read_script_target = read_script_target.replace('\n', ' ')
        read_script_target = read_script_target.replace('.', '. \n')
        read_script_target = read_script_target.replace('?' , '? \n')
        read_script_target = read_script_target.replace('>>' ,'\n >>')
        new_script_target+=' '
        new_script_target+= read_script_target


# In[19]:


new_script_target


# In[20]:


doc = Document()
doc.add_heading(f'{utb_title} YouTube Transcript', level=1)  # 문서 제목 추가
doc.add_paragraph(new_script_target)  # 스크립트 추가
# 문서 저장
doc.save('transcript_v18_only_ko.docx')
print("transcript_v18_only_ko.docx로 저장되었습니다.")


# 
# **제미니 테스트 반복적인 api 문제점 인식이 어렵다**
# 
# 굿굿 만족 무료라면 괜찮네
# 
# 일단 번역 같은 경우는 무료를 쓰고.. 영어 단어 및 구문을 찾아 추가적으로 설명하면 될거 같긴하네
# 
# 이정도면 openai 도 괜찮을 거 같다
# 
# . 이 없는 경우는 추가적으로 한번 돌리고 그럼되나
# 
# . 이 없다면 힘들겠는데
# 
# 추가적으로 리소스 할당이 적다

# In[78]:


import google.generativeai as genai

genai.configure(api_key="AIzaSyAFGo_OBA3Mzuqpa6Qgn8dQGF56zRJOwGk")
model = genai.GenerativeModel("gemini-1.5-flash")
response = model.generate_content("Explain how AI works")

generation_config=genai.types.GenerationConfig(
            # Only one candidate for now.
            candidate_count=1,
            stop_sequences=["x"],
            temperature=0,
        )

#print(response.text) 테스트문장


# In[22]:


chunk_size = 300
chunks_script = [result_transcript[i:i + chunk_size] for i in range(0, len(result_transcript), chunk_size)]
    
response_text =[]

for chunk in chunks_script:
    
   #Please add delimiters (periods and question marks) to the given text.
    response = model.generate_content(
        f"""{chunk} 텍스트 내에 고급 단어가 포함된 경우, 다음 형식으로 단어 설명을 제공해 주세요:

**[문장 내 시간] 단어 : 뜻** - 문맥 설명
   - '시간'은 문장에서 단어가 포함된 시간을 나타내며, '뜻'에는 단어의 의미가 들어갑니다.
   - '문맥 설명'은 해당 단어가 사용된 맥락에서 어떤 의미를 전달하는지를 요약합니다.

예시:
**[09:08] taunt : 상대를 자극하거나 조롱하는 행위** - 특정 인물을 비판하고 자극하는 의미로 사용되었습니다.

이 형식을 참고하여 텍스트 내 고급 단어들의 설명을 작성해 주세요.""" ,
        generation_config=generation_config
        
    )
    response_text.append(response.text)
    response_text.append("\n")
# print(response_text)
print(response.text)

advanced_word = "".join(response_text)


# In[23]:


adw = advanced_word.splitlines()
adw[10]


# In[24]:


doc = Document()
doc.add_heading(f'{utb_title} YouTube Transcript', level=1)   # 문서 제목 추가
doc.add_paragraph(advanced_word)  # 스크립트 추가

# 문서 저장
doc.save('transcript_gemini_kor_v18.docx')
print("transcript_gemini_kor_v18.docx로 저장되었습니다.")


# **자동완성으로 만든 한글말고 번역을 써주는 코드** 

# **facebook/mbart-large-50-many-to-many-mmt 모델 번역 특화 모델**

# In[69]:


from transformers import MBartForConditionalGeneration, MBart50TokenizerFast



def Mbart_model(text):

    model = MBartForConditionalGeneration.from_pretrained("facebook/mbart-large-50-many-to-many-mmt")
    tokenizer = MBart50TokenizerFast.from_pretrained("facebook/mbart-large-50-many-to-many-mmt")

    
    tokenizer.src_lang = "en_XX"
    encoded_hi = tokenizer(text, return_tensors="pt")

    generated_tokens = model.generate(
        **encoded_hi,
        forced_bos_token_id=tokenizer.lang_code_to_id["ko_KR"]
    )
    return tokenizer.batch_decode(generated_tokens, skip_special_tokens=True)

Mbart_model("It says that Ukraine was given a heads up because president-elect, Trump had spoken to Ukraine's Pleasant President Vladimir stalinsky.")


# **구글 번역**

# In[72]:


from translate import Translator

# 번역할 언어 설정
# 예: 영어에서 한국어로 번역
# 13분 짜리 처리 시간 4분 ㄷㄷ 어쩌지 이거 써야하나
def used_translator(text):
    translator = Translator(to_lang="ko", from_lang="en")
    translation_script = translator.translate(text)
    return translation_script
used_translator("It says that Ukraine was given a heads up because president-elect, Trump had spoken to Ukraine's Pleasant President Vladimir stalinsky.")


# **제미니**

# In[82]:


def gemini(text):    
   #Please add delimiters (periods and question marks) to the given text.
    response2 = model.generate_content(
       f"""{text}를 한국어로 번역해주세요""",
        generation_config=generation_config
    )
    return response2.text


gemini("It says that Ukraine was given a heads up because president-elect, Trump had spoken to Ukraine's Pleasant President Vladimir stalinsky.")


# **유사도 기반 한글 자막과 영어 자막의 유사도를 기반으로 파악한다**
# 
# 장점 빠름 
# 부정확한 것이 있을 수도 있다
# 
# Mbart 모델과 구글 자동번역의 결합으로 시간을 단축하였음 
# 
# 8분짜리 영상이 40초안에 끝이 난다.
# 
# 물론 아직도 문제가 있는 부분이 있을것이다.
# 
# 유사도 기반이면 유사도가 낮아도 들어갈수있지 않은가?
# 
# 다른 영상을 테스트 하면서 판단 해야할듯 하다
# 
# 16분 30초로 다운 시켰다

# **mbart - paraphrase-multilingual-MiniLM-L12-v2  
# 
# 
# ** paraphrase-xlm-r-multilingual-v1         
# 0.5 - > 10초
# 어색한 문장 있음 
#                                             
# 

# In[61]:


from sentence_transformers import SentenceTransformer  # 텍스트 백터 변환
from sklearn.metrics.pairwise import cosine_similarity # 벡터 유사도 계산
import numpy as np

# 파일 읽기 리스트화 하였습니다
english_lines = result_transcript

korean_lines = new_script_target.splitlines() 

# 문장 임베딩 모델 로드 (다국어 지원 모델 사용)
#model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')#paraphrase-xlm-r-multilingual-v1
model = SentenceTransformer('paraphrase-xlm-r-multilingual-v1')
# 영어와 한글 문장의 임베딩 벡터 생성 # 임베딩 생성
english_embeddings = model.encode(english_lines)
korean_embeddings = model.encode(korean_lines)

# 유사도 매트릭스 계산
similarity_matrix = cosine_similarity(english_embeddings, korean_embeddings) 

# 유사도가 가장 높은 문장끼리 매칭
merged_lines = []
used_korean_indices = set() # 사용한 한국어는 지우기 위해 집합 사용

for eng_idx, eng_sentence in enumerate(english_lines):
    # 각 영어 문장에 대해 가장 유사한 한글 문장을 찾음
    if eng_sentence == "\n":
        continue
   
    eng_time_rm = re.sub(r'\[\d{2}:\d{2}\]','', eng_sentence) # [] 이거 때문에 번역 잘안나오는듯?
   
    best_kor_idx = np.argmax(similarity_matrix[eng_idx])
    best_kor_similarity = similarity_matrix[eng_idx, best_kor_idx]
    
     
    # 이미 사용된 한글 문장이 아닐 경우에만 매칭 (중복 매칭 방지)
    if best_kor_idx not in used_korean_indices:
        
        if best_kor_similarity < 0.686:
            kor_sentence = Mbart_model(eng_time_rm)
            
        else:
            kor_sentence = korean_lines[best_kor_idx]

        used_korean_indices.add(best_kor_idx)
        # 어쩄든 집합 인덱스에 넣어주었다 부정확한 유사도에 적합한 문장은 
        # 다른 문장에도 문제가 있을 것이라 생각하였다,

    else:
        kor_sentence = Mbart_model(eng_time_rm) 
        

    # 매칭 결과 저장
    merged_lines.append(f"{eng_sentence}\n\n{kor_sentence}\n\n{best_kor_similarity}\n\n")
    
    print(eng_sentence) 
    print("\n")
    print(kor_sentence)

    #merged_lines.append(f"{eng_sentence}\n\n{kor_sentence}\n유사도: {best_kor_similarity:.2f}\n")
merged_en_ko_script = "".join(merged_lines)



# In[63]:
if __name__ == "__main__":

    doc = Document()
    doc.add_heading(f'{utb_title} YouTube Transcript', level=1)  # 문서 제목 추가
    doc.add_paragraph(merged_en_ko_script)  # 스크립트 추가
    # 문서 저장
    doc.save('transcript_v1_0_686.docx')
    print("transcript_v1_0_686.docx로 저장되었습니다.")  

    print("유사도 기반으로 자동 매칭된 영어-한글 문장이 포함된 파일이 생성되었습니다.")


# In[28]:


adw[13]


# In[29]:


Shadowing_script = []

for i in range(len(merged_lines)):
    time_judge = re.search(r"\[(\d{2}:\d{2})\]", merged_lines[i])

    if time_judge:  # time_judge가 None이 아닐 때
        time_str = time_judge.group(0)  # 추출된 시간 문자열 저장
        Shadowing_script.append(merged_lines[i])

        for j in range(len(adw)):
            if time_str in adw[j]:  # time_str이 adw[j]에 있는지 확인
                #Shadowing_script.append(merged_lines[i])
                Shadowing_script.append(adw[j].replace(time_str,""))
                Shadowing_script.append("\n\n")
                #break  # 일치하는 항목을 찾았으므로 중복 추가 방지 위해 반복 종료
            #else:
                #Shadowing_script.append(merged_lines[i])
    else:
        Shadowing_script.append(merged_lines[i])


# In[30]:


Shadowing_script = ["\n\n"]

for i in range(len(merged_lines)):
    time_judge = re.search(r"\[(\d{2}:\d{2})\]", merged_lines[i])

    if time_judge:  # time_judge가 None이 아닐 때
        time_str = time_judge.group(0)  # 추출된 시간 문자열 저장
        
        Shadowing_script.append(merged_lines[i])

        for j in range(len(adw)):
            if time_str in adw[j]:  # time_str이 adw[j]에 있는지 확인
                Shadowing_script.append(adw[j].replace(time_str,""))
                Shadowing_script.append("\n\n")
                    
  


# In[31]:


Shadowing_script_word = "".join(Shadowing_script)


# In[32]:


from docx import Document
from docx.shared import Pt

# 문서 생성
doc = Document()

# 제목 추가
doc.add_heading(f'{utb_title} YouTube Transcript', level=1)

# 스크립트 추가
para = doc.add_paragraph(Shadowing_script_word)

# 문단 내 모든 텍스트의 폰트 크기 변경
for run in para.runs:
    run.font.size = Pt(12)  # 폰트 크기를 12포인트로 설정

# 문서 저장
doc.save('transcript_v18_ko)en-word.docx')
print("transcript_v18_ko)en_word.docx로 저장되었습니다.")  

print("유사도 기반으로 자동 매칭된 영어-한글 문장이 포함된 파일이 생성되었습니다.")


# In[33]:




# **개발 완료 -영어 한국어 고급단어** 

# **이로서 어휘 빼면 모든 개발이 끝난 것 같다**

# **facebook/mbart-large-50-many-to-many-mmt 모델 번역 특화 모델**
# 
# 이것만 사용했을 경우 시간이 얼마나 소요될까
# 
# 8분짜리  14분 소요 
# 
# 꽤 걸리는 모습을 보여준다 그래도 엄청 걸린다는 아닌거 같다
# 
# 하지만 해당 모델을 1시간 걸리는 인터뷰라든지 쓰는 것은 부적절해 보인다

# In[34]:


# # tesst = ["\n"]

# # for i in result_transcript:
    
# #     if i =="\n":
# #         print("널")
# #         continue

# #     MM = Mbart_model(i)    
# #     tesst.append(i)
# #     tesst.append(i)
# #     tesst.append(MM)
# #     tesst.append("\n")

# #     print(i)
# #     print("\n")
# #     print(MM)
# #     print("\n")


# **Duration 기반으로 en 에 .이 없다면 

# **from translate import Translator**
# 
# 무료아님... 나중에는 돈내야함
# 
# 잘되긴했다 하지만 무료가 아니었음.. 
# 근데 누가 서버비까지 내어주면서 이걸 할까 ㅋㅋ

# In[32]:


# from translate import Translator

# # 번역할 언어 설정
# # 예: 영어에서 한국어로 번역
# # 13분 짜리 처리 시간 4분 ㄷㄷ 어쩌지 이거 써야하나

# translator = Translator(to_lang="ko", from_lang="en")

# # 번역할 문장
# new_script_line = new_script.splitlines()

# result_transcript = " "

# for script_line in new_script_line:

#     # 번역 실행
#     translation_script = translator.translate(script_line)

#     # 결과 출력
#     result_transcript += script_line
#     result_transcript += "\n\n"
#     result_transcript +=  translation_script
#     result_transcript += "\n\n"


# In[33]:


# doc = Document()
# doc.add_heading('YouTube Transcript', level=1)  # 문서 제목 추가
# doc.add_paragraph(result_transcript)  # 스크립트 추가

# # 문서 저장
# doc.save('transcript_v30_translate.docx')
# print("transcript_v30_translate.docx로 저장되었습니다.")


# **유튜브API 분석 모르는게 있나해서**

# In[ ]:


# from youtube_transcript_api import YouTubeTranscriptApi

# # retrieve the available transcripts
# transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

# # iterate over all available transcripts
# for transcript in transcript_list:

#     # the Transcript object provides metadata properties
#     print(
#         transcript.video_id,
#         transcript.language,
#         transcript.language_code,
#         # whether it has been manually created or generated by YouTube
#         transcript.is_generated,
#         # whether this transcript can be translated or not
#         transcript.is_translatable,
#         # a list of languages the transcript can be translated to
#         transcript.translation_languages,
#     )

#     # fetch the actual transcript data
#     print(transcript.fetch())

#     # translating the transcript will return another transcript object
#     print(transcript.translate('en').fetch())

# # you can also directly filter for the language you are looking for, using the transcript list
# transcript = transcript_list.find_transcript(['de', 'en'])  

# # or just filter for manually created transcripts  
# transcript = transcript_list.find_manually_created_transcript(['de', 'en'])  

# # or automatically generated ones  
# transcript = transcript_list.find_generated_transcript(['de', 'en'])


# In[ ]:





# In[ ]:





# **아래 실패한 것들**
# 
# 정리
# llama 모델을 돌려서 번역과 고급 단어 분석
#     시간 과다 소요 및 할루시 , 명령에 맞지 않는 말 추가
#     -> 병렬처리 X 시간 약간 줄지만 그래도 길었다
#     -> 양자화로 모델을 경량화 시간 약간 줄지만 길었으며 명령에 맞지 않는 말
#     ->명령에 맞지 않는 말을 고치기 위해 프롬프트 개선 및 langchain  에서 facebook으로 변경 
#     많이 좋아졌지만 없진 않았음 
#     eeve10B 보다 작고 성능은 최신인 llama 모델 사용 시간 약간 줄었지만
#     시간 과다 소요를 고치지 못해 
#     사용 불가 결정하였다.
#     공부를 위해 코드는 남겨두었음
# 

# 모델 경량화 기법은 대규모 모델의 크기를 줄이거나 성능을 유지하면서 효율성을 높이기 위해 다양한 방법을 사용합니다. 여기 몇 가지 주요 경량화 기법과 그 특징을 설명할게요:
# 
# 1. 지식 증류 (Knowledge Distillation)
# 개념: 큰 모델(교사 모델)에서 학습한 지식을 작은 모델(학생 모델)에 전이하는 기법입니다. 학생 모델은 교사 모델의 예측을 기반으로 학습하여 성능을 유지하면서 더 작은 크기로 만들어집니다.
# 장점: 모델 크기가 작아지면서도 높은 성능을 유지할 수 있습니다.
# 2. 모델 프루닝 (Model Pruning)
# 개념: 중요하지 않거나 기여도가 낮은 파라미터를 제거하여 모델을 경량화하는 기법입니다. 이를 통해 연산량과 메모리 사용량을 줄일 수 있습니다.
# 장점: 경량화가 진행되며, 일반적으로 모델의 성능 저하가 최소화됩니다.
# 3. 양자화 (Quantization)
# 개념: 모델의 파라미터와 연산을 저비트 정수로 변환하여 메모리와 계산 효율성을 높이는 기법입니다. 예를 들어, 32비트 부동 소수점 대신 8비트 정수로 표현할 수 있습니다.
# 장점: 모델 크기를 크게 줄일 수 있으며, 특히 하드웨어에서의 계산 속도를 높이는 데 효과적입니다.
# 4. 모델 압축 (Model Compression)
# 개념: 다양한 압축 기법을 사용하여 모델의 전체 크기를 줄이는 방법입니다. 이에는 지식 증류, 프루닝, 양자화 등이 포함될 수 있습니다.
# 장점: 다양한 기술을 조합하여 최적의 성능과 효율성을 얻을 수 있습니다.
# 5. 신경망 아키텍처 최적화
# 개념: 더 경량화된 아키텍처(예: MobileNet, SqueezeNet)를 사용하여 자연스럽게 경량화된 모델을 만드는 것입니다. 이러한 아키텍처는 효율적인 연산을 위해 설계되었습니다.
# 장점: 처음부터 경량화된 구조로 설계되기 때문에 성능과 효율성의 균형이 잘 맞습니다.
# 6. 하이퍼파라미터 조정
# 개념: 모델의 하이퍼파라미터(예: 레이어 수, 노드 수)를 조정하여 필요한 성능을 유지하면서도 크기를 줄이는 방법입니다.
# 장점: 모델의 복잡성을 조절하여 효율성을 높일 수 있습니다.
# 7. 전이 학습 (Transfer Learning)
# 개념: 대규모 데이터셋에서 훈련된 모델을 사용하여 특정 작업에 대한 적응을 빠르게 수행할 수 있도록 하는 방법입니다. 이 과정에서 필요 없는 레이어를 제거하거나 조정하여 경량화할 수 있습니다.
# 장점: 빠른 훈련과 성능 유지를 통해 효율성을 높일 수 있습니다.
# 결론
# 각 경량화 기법은 특정 상황과 모델에 따라 다르게 적용될 수 있으며, 최적의 경량화 방법은 사용자의 필요와 목표에 따라 달라질 수 있습니다. 여러 기법을 조합하여 최적의 성능과 효율성을 찾아내는 것이 중요합니다.

# In[2]:


# import os 
# os.environ['TOKEN'] = "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"


# In[1]:


# from transformers import AutoTokenizer , AutoModelForCausalLM
# import torch

# model_id = "Bllossom/llama-3.2-Korean-Bllossom-3B"

# tokenizer = AutoTokenizer.from_pretrained(model_id)

# model = AutoModelForCausalLM.from_pretrained(
#     model_id,
#     torch_dtype = torch.bfloat16,
#     device_map = "auto",
# )


# Hugging 을 이용한 eeve-korean 모델 사용 (야놀자)
# 랭서버로 구동 실습 및 프롬프트 생성
# 
# 하지만 답변은 잘하지만 문제가 있긴하다 
# 
# 정해진 명령을 수행을 잘못함
# 
# * 영어 단어 테스트 수행은 괜찮음 만족 
# 
# 번역기 돌린 후 집어 넣고 구글 넣고 여기에 있다면 추출하는 방식이 좋을거 같다
# openai 가 성능은 더 좋겠지 아마 
# 운영비 줄일수 있겠는데 성능은 모르겠고
# 
# 

# 중간에 이상한 말들이 끼네
# 사용불가함 

# **chatollama 형식 및 langchain 으로 작성한 코드**
# 
# 전반적으로 쓸데 없는말이 끼어있는경우가 많다 
# 
# 모델 자체 성능이 구린건지 아님 형식 입력이 명확하게 들어가지 않는건지 파악하기 어렵다

# In[ ]:


# from langchain_community.chat_models import ChatOllama
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
# from langchain_core.callbacks.manager import CallbackManager

# # LangChain이 지원하는 다른 채팅 모델을 사용합니다. 여기서는 Ollama를 사용합니다.
# llm = ChatOllama(
#     model="Bllosom_llama_3.2:latest ",
#     temperature=0,
#     callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
# )


# prompt = ChatPromptTemplate.from_template(
#     "너는 번역을 수행하는 챗봇이야 , 다음 내용을 한국어로 번역해줘 {topic}.\n",
# )

# # LangChain 표현식 언어 체인 구문을 사용합니다.
# chain = prompt | llm | StrOutputParser()

# # 번역할 문장
# new_script_line = new_script.splitlines()

# result_transcript = " "

# for script_line in new_script_line[2:]:

#     # 결과 출력
#     result_transcript += script_line
#     result_transcript += "\n\n"
#     result_transcript += chain.invoke({"topic" : script_line}) # 번역 실행 


# In[ ]:


# result_transcript


# <!-- ** -->

# In[10]:


# import torch
# from transformers import BitsAndBytesConfig

# bnb_config = BitsAndBytesConfig(
  #  load_in_4bit=True
# )


# In[11]:


# import torch
# from transformers import AutoTokenizer, AutoModelForCausalLM
# import bitsandbytes as bnb

# model_id = 'ai-human-lab/EEVE-Korean_Instruct-10.8B-expo'

# tokenizer = AutoTokenizer.from_pretrained(model_id)
# model = AutoModelForCausalLM.from_pretrained(
#     model_id,
#     quantization_config = bnb_config,
#     torch_dtype=torch.bfloat16,
    
# )

# # 모델을 특정 GPU로 이동시킵니다.
#   # 기본 GPU 장치로 이동

# def generate_response(system_message , user_message):

#     messages = [
#         {"role": "system", "content": system_message},
#           {"role": "user", "content": user_message}
#         ]

#     input_ids = tokenizer.apply_chat_template(
#         messages,
#         add_generation_prompt=True,
#         return_tensors="pt"
#     ).to(model.device)


#     terminators = [
#         tokenizer.eos_token_id,
#         tokenizer.convert_tokens_to_ids("<|eot_id|>")
#     ]

#     outputs = model.generate(
#         input_ids,
#         max_new_tokens=256,
#         eos_token_id= terminators,
#         do_sample=True,
#         temperature=0.6,
#         top_p=0.9,
#     )

#     return(tokenizer.decode(outputs[0][input_ids.shape[-1]:], skip_special_tokens=True))


# In[ ]:


# from langchain_community.chat_models import ChatOllama
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
# from langchain_core.callbacks.manager import CallbackManager

# # 번역할 문장 13분짜리 영상이 28분 걸렸음 번역작업 수행은 만족 
# #  
# new_script_line = new_script.splitlines()

# result_transcript = " "

# # llama3_translation_text=generate_response(
# #         system_message= "너는 번역을 수행하는 챗봇이야 다음 내용을 한국어로 번역해줘",
# #                                             user_message="The sun was setting behind the mountains, painting the sky with shades of orange and pink.") 

# # for script_line in new_script_line[2:]:

# #     # 결과 출력
# #     llama3_translation_text=generate_response(
# #         system_message= "너는 번역을 수행하는 챗봇이야 다음 내용을 한국어로 번역해줘",
# #                                             user_message=script_line) 

# #     result_transcript += script_line
# #     result_transcript += "\n\n"
# #     result_transcript +=  llama3_translation_text# 번역 실행 

# lines = []
# for script_line in new_script_line[2:]:
#     llama3_translation_text = generate_response(
#         system_message="너는 번역을 수행하는 챗봇이야. 사용자가 제공하는 텍스트를 정확하게 한국어로 번역해야 해. 번역은 자연스럽고 명확해야 하며, 의미가 왜곡되거나 추가 설명을 하지 않도록 해. 각 단어의 의미를 그대로 전달해야 해.",
#         user_message=script_line
#     )
    
#     # 원본 스크립트와 번역을 리스트에 추가
#     lines.append(script_line)
#     lines.append("\n\n")
#     lines.append(llama3_translation_text)
#     print(llama3_translation_text)
#     lines.append("\n\n")
    
# # 마지막에 한 번에 합침
# result_transcript = ''.join(lines)


# In[30]:


#print(result_transcript)


# In[13]:


# def generate_batch_response(system_message, user_messages):
#     messages = [{"role": "system", "content": system_message}]
#     user_inputs = [{"role": "user", "content": msg} for msg in user_messages]
#     messages.extend(user_inputs)

#     input_ids = tokenizer.apply_chat_template(
#         messages,
#         add_generation_prompt=True,
#         return_tensors="pt"
#     ).to(model.device)

#     terminators = [
#         tokenizer.eos_token_id,
#         tokenizer.convert_tokens_to_ids("<|eot_id|>")
#     ]

#     outputs = model.generate(
#         input_ids,
#         max_new_tokens=256,
#         eos_token_id=terminators,
#         do_sample=True,
#         temperature=0.6,
#         top_p=0.9,
#     )

#     # 여러 줄의 응답을 반환
#     return [tokenizer.decode(output[input_ids.shape[-1]:], skip_special_tokens=True) for output in outputs]

# new_script_line = new_script.splitlines()

# result_transcript = " "


# # 여러 줄을 한 번에 처리
# user_messages_batch = new_script_line[2:]  # 예시로 전체 스크립트 라인을 배치로 처리
# translations = generate_batch_response(
#     system_message="너는 번역을 수행하는 챗봇이야. 사용자가 제공하는 텍스트를 정확하게 한국어로 번역해야 해.",
#     user_messages=user_messages_batch
# )

# # 응답 처리
# result_transcript = '\n\n'.join([line + "\n\n" + translation for line, translation in zip(new_script_line[2:], translations)])


# In[15]:


#print(result_transcript)


# In[ ]:


# import concurrent.futures

# def generate_response_parallel(script_line):
#     return generate_response(
#         system_message="너는 번역을 수행하는 챗봇이야. 사용자가 제공하는 텍스트를 정확하게 한국어로 번역해야 해.",
#         user_message=script_line
#     )

# # 병렬 처리로 번역
# with concurrent.futures.ThreadPoolExecutor() as executor:
#     translations = list(executor.map(generate_response_parallel, new_script_line[2:]))

# # 결과 조합
# result_transcript = '\n\n'.join([line + "\n\n" + translation for line, translation in zip(new_script_line[2:], translations)])


# In[ ]:





# In[ ]:





# In[31]:


# doc = Document()
# doc.add_heading(f'{utb_title} YouTube Transcript', level=1)  # 문서 제목 추가
# doc.add_paragraph(result_transcript)  # 스크립트 추가

# # 문서 저장
# doc.save('transcript_v35_t.docx')
# print("transcript_v35_t.docx로 저장되었습니다.")


# In[ ]:


# from langchain_community.chat_models import ChatOllama
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
# from langchain_core.callbacks.manager import CallbackManager

# # LangChain이 지원하는 다른 채팅 모델을 사용합니다. 여기서는 Ollama를 사용합니다.
# llm = ChatOllama(
#     model="llama-3.2-Korean-Bllossom-3B-gguf-Q4_K_M",
#     temperature=0,  # 응답의 창의성 조절
#      # 최대 토큰 수
#     callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
# )
# prompt = ChatPromptTemplate.from_template(   "다음 주제에 대해 TOEIC 850 수준의 영어 단어를 나열하고, 각 단어의 뜻을 다음 형식으로 작성해 주세요:\n"
#     "- 단어: [단어]\n"
#     "  뜻: [뜻]\n"
#     "주제: {topic}")

# # LangChain 표현식 언어 체인 구문을 사용합니다.
# chain = prompt | llm | StrOutputParser()

# # 간결성을 위해 응답은 터미널에 출력됩니다.

# chunk_size = 300
# modified_script = [new_script[i:i + chunk_size] for i in range(0, len(new_script), chunk_size)]
# hard_word =" "

# for modified_script_split in modified_script:
#     hard_word += chain.invoke({"topic":    modified_script_split})




# In[ ]:


#modified_script


# In[3]:


# from langchain_community.chat_models import ChatOllama
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
# from langchain_core.callbacks.manager import CallbackManager

# # LangChain이 지원하는 다른 채팅 모델을 사용합니다. 여기서는 Ollama를 사용합니다.
# llm = ChatOllama(
#     model="Bllossom/llama-3.2-Korean-Bllossom-3B",
#     temperature=0,  # 응답의 창의성 조절
#      # 최대 토큰 수
#     callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
# )
# prompt = ChatPromptTemplate.from_template(   "다음 주제에 대해 TOEIC 850 수준의 영어 단어를 나열하고, 각 단어의 뜻을 다음 형식으로 작성해 주세요:\n"
#     "- 단어: [단어]\n"
#     "  뜻: [뜻]\n"
#     "주제: {topic}")

# # LangChain 표현식 언어 체인 구문을 사용합니다.
# chain = prompt | llm | StrOutputParser()

# # 간결성을 위해 응답은 터미널에 출력됩니다.

# chunk_size = 300
# modified_script = [new_script[i:i + chunk_size] for i in range(0, len(new_script), chunk_size)]
# hard_word =" "

# for modified_script_split in modified_script:
#     hard_word += chain.invoke({"topic":    modified_script_split})




# In[35]:





# In[ ]:


# from youtube_transcript_api import YouTubeTranscriptApi
# from docx import Document

# new_script2 = "\n\n"

# modified_script= modified_script.replace('.' , '. \n\n')
# modified_script= modified_script.replace('?' , '? \n\n')
# modified_script= modified_script.replace('>>', '\n\n >>')

# new_script2 += modified_script

# doc = Document()
# doc.add_heading('YouTube Transcript', level=1)  # 문서 제목 추가
# doc.add_paragraph(new_script2)  # 스크립트 추가

# # 문서 저장
# doc.save('transcript_v21_kor.docx')
# print("transcript_v21_kor.docx로 저장되었습니다.")


# In[ ]:





# In[ ]:





# In[ ]:


# #langserver  을 통한 
# from langserve import RemoteRunnable

# # ngrok remote 주소 설정

# #chain = RemoteRunnable("http://127.0.0.1:8000/prompt/c/N4XyA/")
# chain = RemoteRunnable(" https://57c7-1-225-116-182.ngrok-free.app/prompt/")
# # chain = RemoteRunnable("http://0.0.0.0:8000/prompt/")

# for token in chain.stream({"topic": "Openai has wowed users with the  Chat bot's thoughtful answers,  The ability to create new  Content, but couldn't tell you  About things like the yankees. 해당 문장을 한국어로 번역하고 토익 850 수준의 영단어가 있으면 간단하게 기재하고 번역해 "}):
#     print(token, end="")


# LLM 을 Runnable로 실행

# In[ ]:


# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from langserve import RemoteRunnable
# # llm.py llm 라우팅 llm 주소에 접근하면 내 모델을 가져간다 
# # llm = RemoteRunnable("https://poodle-deep-marmot.ngrok-free.app/llm/")
# llm = RemoteRunnable("https://57c7-1-225-116-182.ngrok-free.app/llm/")

# prompt = ChatPromptTemplate.from_template(
#     "다음의 내용을 번역하고 토익 850점 수준의 영어단어를 기재해주세요:\n{input}"
# )


# chain = prompt | llm | StrOutputParser()


# In[ ]:


# from langchain_openai import ChatOpenAI

# llama3 = ChatOpenAI(
#     base_url="",
#     api_key=,  
#     model=,
#     temperature=0.1
# )

