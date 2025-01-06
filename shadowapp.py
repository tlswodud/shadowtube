import streamlit as st

#yt_dlp 새해 기념으로 유튜브에서 밴하였음
#쿠키로 허용가능하게 할수있지만 최근 전부 사용불가해졌음 유튜브 api 를 통해 변경
import yt_dlp
import streamlit as st
from io import BytesIO
from docx import Document

import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from urllib.parse import urlparse, parse_qs
import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi


from docx import Document
from docx.shared import Pt
from io import BytesIO

from docx.oxml import OxmlElement
# def get_video_info(url):
#     try:
#         # # with open('cookies.json', 'r') as f:
#         # #     cookies = json.load(f)
#         ydl_opts = {
#             'quiet': True,
#             'no_warnings': True,
#             'extract_flat': True,
#             'skip_download': True,  # 다운로드 건너뛰기
#             'prefer_ffmpeg': False,
            

#         }
    
#         with yt_dlp.YoutubeDL(ydl_opts) as ydl:
#             info = ydl.extract_info(url, download=False)
#             return {
#                 'title': info.get('title'),
#                 'channel': info.get('uploader'),
#                 'thumbnail': info.get('thumbnail'),
#                 'duration': info.get('duration'),
#                 'view_count': info.get('view_count')
#             }
#     except Exception as e:
#         st.error(f"비디오 정보를 가져오는데 실패했습니다: {str(e)}")
#         return None
import requests    
@st.cache_data
def get_video_id(url):
    """URL에서 유튜브 비디오 ID 추출"""
    parsed_url = urlparse(url)
    
    # 일반 유튜브 URL (https://www.youtube.com/watch?v=비디오ID)
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path == '/watch':
        return parse_qs(parsed_url.query).get('v', [None])[0]
    
    # 짧은 URL 형식 (https://youtu.be/비디오ID)
    if parsed_url.hostname == 'youtu.be':
        return parsed_url.path[1:]
    
    # 임베드 URL 형식 (https://www.youtube.com/embed/비디오ID)
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path.startswith('/embed/'):
        return parsed_url.path.split('/')[2]
    
    # 유튜브 쇼츠 URL 형식 (https://www.youtube.com/shorts/비디오ID)
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path.startswith('/shorts/'):
        return parsed_url.path.split('/')[2]
    
    return None

def get_video_info(url, youtube_api_key_):
    # URL에서 비디오 ID 추출
    video_id = get_video_id(url)  # 'v=' 뒤에 오는 ID를 추출하는 방법
    # API 요청 URL 구성
    api_url = f"https://www.googleapis.com/youtube/v3/videos?part=snippet,statistics,contentDetails&id={video_id}&key={youtube_api_key_}"
    
    # API 요청
    response = requests.get(api_url)
    
    # 응답이 정상적인지 확인
    if response.status_code == 200:
        video_info = response.json()
        
        # 'items' 리스트가 비어있지 않은지 확인
        if 'items' in video_info and len(video_info['items']) > 0:
            # 통계 정보와 스니펫 정보, 콘텐츠 세부 정보 추출
            statistics = video_info['items'][0].get('statistics', {})
            snippet = video_info['items'][0].get('snippet', {})
            content_details = video_info['items'][0].get('contentDetails', {})
            
            # duration 정보가 없을 경우 기본값 처리
            duration = content_details.get('duration', '정보 없음')
            
            # 필요한 정보 반환
            return {
                'title': snippet.get('title'),
                'channel': snippet.get('channelTitle'),
                'description': snippet.get('description'),
                'thumbnail': snippet.get('thumbnails', {}).get('high', {}).get('url', None),
                'statistics': statistics,
                #'duration': duration
            }
        else:
            print("비디오 정보를 찾을 수 없습니다.")
            return None
    else:
        print(f"API 요청 실패: {response.status_code}")
        return None
    


import base64
def create_modern_ui():
    # 헤더 섹션

    st.markdown("""
    <head>
        <meta name="google-adsense-account" content="ca-pub-4095011834932682">                     
    </head>
    """, unsafe_allow_html=True)
    st.markdown("""
    <script async src="https://pagead2.googlesyndication.com/pagead/js/adsbygoogle.js?client=ca-pub-4095011834932682"
     crossorigin="anonymous"></script>
    """, unsafe_allow_html=True)
    
    with open("./image/shadowLogo2.png", "rb") as image_file:
         encoded_Logo = base64.b64encode(image_file.read()).decode()
                   
    st.markdown(f"""
    <style>
        @media (max-width: 600px) {{
            .shadowtube-title {{
                font-size: 1.5rem; /* 핸드폰에서의 글씨 크기 */
            }}
        }}
    </style>
    <div style="display: flex; justify-content: center; align-items: center; height: 20vh; flex-direction: column;">
            <div style='display:flex;align-items: center; padding: 1rem 0;'>
                <img src="data:image/png;base64,{encoded_Logo}" width="47">
                &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;
                <h1 class='shadowtube-title' style='color: #1E88E5;'>ShadowTube</h1>
            </div>
        <p style='font-size: 1.2rem; color: #424242;'>AI Shadowing Script Generator</p>
    </div>
    """, unsafe_allow_html=True)

    # 소개 섹션 - 카드 스타일 디자인
    st.markdown("""
        <div style='background-color: #F8F9FA; padding: 2rem; border-radius: 10px; margin: 1rem; max-width: 100%; overflow: hidden;'>
            <h3 style='color: black;'>🚀 Learn Languages Through YouTube!</h3>
            <p style='font-size: 1.1rem; color: #424242;'>
                Easily transform your favorite YouTube videos into powerful learning materials!
            </p>
            <div style='display: flex; gap: 1rem; margin-top: 1rem; flex-wrap: wrap;'>
                <div style='background: #E3F2FD; padding: 1rem; border-radius: 8px; flex: 1; min-width: 200px;color: #000;'>
                    <h4>✨ Features</h4>
                    <ul style='margin: 0;'>
                        <li>Translations</li>
                        <li>Vocabulary analysis</li>
                        <li>Shadowing scripts</li>
                    </ul>
                </div>
                <div style='background: #E8F5E9; padding: 1rem; border-radius: 8px; flex: 1; min-width: 200px;color: #000;'>
                    <h4>📚 Benefits</h4>
                    <ul style='margin: 0;'>
                        <li>Learn from your favorite content.</li>
                        <li>Natural pronunciation</li>
                        <li>Contextual learning</li>
                    </ul>
                </div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    st.divider()
    
    # 언어 선택 섹션 - 2개의 컬럼으로 구성
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 👤 Native Language")
        native_language = st.selectbox(
            "Select your native language",
            ["한국어", "English", "日本語", "中文", "Español", "Français"],
            label_visibility="collapsed"
        )
        
    with col2:
        st.markdown("### 🎯 Target Language")
        target_language = st.selectbox(
            "Select language you want to learn",
            ["English", "한국어", "日本語", "中文", "Español", "Français"],
            label_visibility="collapsed"
        )
    
    # 선택된 언어 정보 표시
    with st.container():
        st.info(f"""
        📍 Selected Configuration:
        - Native Language: {native_language} (code: {get_language_code(native_language)})
        - Target Language: {target_language} (code: {get_language_code(target_language)})
        """)
    
   

    # URL 입력 섹션
    st.markdown("### 🎥 YouTube Video")
    url = st.text_input(
        "Enter YouTube URL",
        placeholder="https://www.youtube.com/watch?v=...",
        label_visibility="collapsed"
    )
    # video_id 추출
   
    
    # 비디오 정보 표시
    if url:
        youtube_api_key = "AIzaSyD0UwH5A3yUn50npxpHLkM7VEuaC6GRlqw"
        video_info = get_video_info(url,youtube_api_key)
        
        if video_info:
            with st.container():
                col1, col2 = st.columns([1, 3])
                with col1:
                    st.image(video_info['thumbnail'], width=160)
                with col2:
                    st.markdown(f"#### {video_info['title']}")
                    st.caption(f"📺 {video_info['channel']}")
                    # 추가 정보 표시
                    # duration_min = video_info['duration'] // 60
                    # duration_sec = video_info['duration'] % 60
                 
                   # st.caption(f"⏱️ {duration_min}:{duration_sec:02d} | 👀 {video_info['view_count']:,} views")
        
    # 처리 버튼
    if st.button("🎯 Generate Shadowing Materials", type="primary", use_container_width=True):
        with st.spinner("Processing your request..."):
            # 처리 로직
            pass

    # 도움말 섹션
    with st.expander("ℹ️ How to use"):
        st.markdown("""
        1. Paste your free Gemini API Key.            
        2. Select your native language
        3. Choose the language you want to learn
        4. Paste a YouTube URL
        5. Click 'Generate' to create your learning materials     
        """)
    return url ,  native_language  ,target_language

def get_language_code(language):
    # 언어 코드 매핑
    codes = {
        "한국어": "ko",
        "English": "en",
        "日本語": "ja",
        "中文": "zh-Hans",
        "Español": "es",
        "Français": "fr"
    }
    return codes.get(language, "unknown")

st.set_page_config(
    page_title="ShadowTube",
    page_icon="./image/shadowLogo.png",
    layout="centered",
   
)
# 스타일링
st.markdown("""
    <style>
    .stButton>button {
        font-size: 1.2rem;
        margin-top: 2rem;
        margin-bottom: 2rem;
    }
    .stSelectbox {
        margin-bottom: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)

# UI 실행
url , native_language  ,want_language = create_modern_ui()

native_code=get_language_code(native_language)

learn_code = get_language_code(want_language)

st.markdown("""
    <style>
    .chat-message {
        padding: 1.5rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
        display: flex;
    }
    .chat-message.user {
        background-color: #f0f2f6;
    }
    .chat-message.assistant {
        background-color: white;
        border: 1px solid #e0e0e0;
    }
    .chat-icon {
        width: 40px;
        margin-right: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)

# 채팅 메시지 표시 함수
def display_chat_message(role, content):
    icon = "👤" if role == "user" else "🤖"
    st.markdown(f"""
        <div class="chat-message {role}">
            <div class="chat-icon">{icon}</div>
            <div>{content}</div>
        </div>
    """, unsafe_allow_html=True)
  


import os
for key, value in os.environ.items():
    print(f"{key}: {value}")



def create_word_file_shadow_script(content, utb_title, learn_code, want_font, native_font, font_size): 
    """
    각 언어에 맞는 폰트를 설정하여 워드 파일 생성
    content: 텍스트 내용
    utb_title: 유튜브 제목
    learn_code: 학습하고자 하는 언어 코드
    native_font: 모국어 폰트
    want_font: 학습 언어 폰트
    font_size: 폰트 크기
    """
    # 워드 문서 객체 생성
    doc = Document()
    doc.add_heading(f'{utb_title} YouTube Transcript', level=1)

    # 유니코드 범위 설정
    language_ranges = {
        "zh": [(0x4E00, 0x9FFF)],  # CJK 통합 한자
        "ko": [(0xAC00, 0xD7AF)],  # 한글
        "ja": [(0x3040, 0x309F), (0x30A0, 0x30FF)],  # 히라가나 + 가타카나
        "en": [(0x0041, 0x005A), (0x0061, 0x007A)],  # 영어 대문자 + 소문자
        "fr": [(0x00C0, 0x017F)],  # 프랑스어
        "es": [(0x00C0, 0x017F)]   # 스페인어
    }

    def is_in_range(char, ranges):
        """주어진 문자(char)가 특정 언어 범위(ranges)에 속하는지 확인"""
        code = ord(char)
        return any(start <= code <= end for start, end in ranges)

    def get_dominant_language(text):
        """주어진 텍스트에서 가장 빈도가 높은 언어를 반환"""
        counts = {lang: 0 for lang in language_ranges.keys()}
        for char in text:
            for lang, ranges in language_ranges.items():
                if is_in_range(char, ranges):
                    counts[lang] += 1
                    break
        return max(counts, key=counts.get)  # 가장 빈도가 높은 언어 반환

    # 각 줄에 대해 처리
    for line in content:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
    

        # 주요 언어 확인
        dominant_language = get_dominant_language(line)

        # 학습 언어와 동일하면 학습 폰트, 아니면 모국어 폰트 적용
        
        if dominant_language == learn_code:
             font_name = want_font
        else:
             font_name = native_font     

        # 폰트 크기와 폰트 설정
        run.font.size = Pt(font_size)
        rpr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")

        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)

        # CJK 문자용 폰트 설정 (동아시아 폰트)
        if dominant_language in ["zh", "ko", "ja"]:
            rFonts.set(qn("w:eastAsia"), font_name)

        rpr.append(rFonts)

    # BytesIO를 사용해 워드 파일을 메모리에 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
# 사용자 입력
import streamlit as st
@st.cache_data
def get_best_english_transcript(video_id,_transcript_list):
        # 선호하는 영어 자막 코드 목록

        # 수동이 더 좋다? 영어는 또 그렇지 않았다
         #find_generated_transcript < - >  find_manually_created_transcript 로 변환 
        #뉴스 방송사들은 수동자막이 더 이상함 또한 자동 자막의 퀄이 영어는 너무 좋다

        english_codes = [ 'en','en-US', 'en-GB', 'en-CA', 'en-AU', 'a.en']
        
        try:
        
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도)
            try:
                transcript = _transcript_list.find_generated_transcript(english_codes)
                print(f"수동 생성된 영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                result = [(entry['start'], entry['text']) for entry in transcript.fetch()]
                return result
            
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript = _transcript_list.find_manually_created_transcript(english_codes)
                    print(f"자동 생성된 영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                    result = [(entry['start'], entry['text']) for entry in transcript.fetch()]
                    return result
                    
                except:
                    # 3. 마지막으로 'en'으로 시작하는 모든 코드 확인
                    available_transcripts = list(transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith('en'):
                            print(f"영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                        result = [(entry['start'], entry['text']) for entry in transcript.fetch()]
                        return result
                    
                    print("영어 자막을 찾을 수 없습니다.")
                    return None
                    
        except Exception as e:
            print(f"자막을 가져오는 중 오류가 발생했습니다: {str(e)}")
            return None
@st.cache_data        
def get_best_english_encode(video_id,_transcript_list):
        

        # 선호하는 영어 자막 코드 목록
        english_codes = ['en-US', 'en-GB', 'en-CA', 'en-AU', 'en', 'a.en']
        
        try:
            #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도)
            try:
                transcript = _transcript_list.find_manually_created_transcript(english_codes)
                print(f"수동 생성된 영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                return english_codes
                
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript = _transcript_list.find_generated_transcript(english_codes)
                    print(f"자동 생성된 영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                    return english_codes
                    
                except:
                    # 3. 마지막으로 'en'으로 시작하는 모든 코드 확인
                    available_transcripts = list(_transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith('en'):
                            print(f"영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                            return english_codes
                    
                    print("영어 자막을 찾을 수 없습니다.")
                    return None
                    
        except Exception as e:
            print(f"자막을 가져오는 중 오류가 발생했습니다: {str(e)}")
            return None
        
#fetch 아 힘들다..
@st.cache_data
def get_best_english_transcript_no_time(video_id,_transcript_list):
        # 선호하는 영어 자막 코드 목록
            english_codes = ['en','en-US', 'en-GB', 'en-CA', 'en-AU', 'a.en']
                    
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도) - >
            # 수동이 더 좋다? 영어는 또 그렇지 않았다
            try:
                #find_generated_transcript < - >  find_manually_created_transcript 로 변환 
                #뉴스 방송사들은 수동자막이 더 이상함 

                transcript = _transcript_list.find_generated_transcript(english_codes)
                print(f"수동 생성된 영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                result = [(entry['text']) for entry in transcript.fetch()]
                return result
            
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript = _transcript_list.find_manually_created_transcript(english_codes)
                    print(f"자동 생성된 영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                    result = [(entry['text']) for entry in transcript.fetch()]
                    return result
                    
                except:
                    # 3. 마지막으로 'en'으로 시작하는 모든 코드 확인
                    available_transcripts = list(_transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith('en'):
                            print(f"영어 자막을 찾았습니다. (언어 코드: {transcript.language_code})")
                        result = [(entry['text']) for entry in transcript.fetch()]
                        return result
                    
                    print("영어 자막을 찾을 수 없습니다.")
                    return None
                    
    


@st.cache_data
def get_video_id(url):
    """URL에서 유튜브 비디오 ID 추출"""
    parsed_url = urlparse(url)
    
    # 일반 유튜브 URL (https://www.youtube.com/watch?v=비디오ID)
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path == '/watch':
        return parse_qs(parsed_url.query).get('v', [None])[0]
    
    # 짧은 URL 형식 (https://youtu.be/비디오ID)
    if parsed_url.hostname == 'youtu.be':
        return parsed_url.path[1:]
    
    # 임베드 URL 형식 (https://www.youtube.com/embed/비디오ID)
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path.startswith('/embed/'):
        return parsed_url.path.split('/')[2]
    
    # 유튜브 쇼츠 URL 형식 (https://www.youtube.com/shorts/비디오ID)
    if parsed_url.hostname == 'www.youtube.com' and parsed_url.path.startswith('/shorts/'):
        return parsed_url.path.split('/')[2]
    
    return None
# 사용자 입력 받기

def get_video_title(video_id):
    # yt-dlp 객체 생성
    if video_id == None:
         return None

    ydl_opts = {}
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        # 비디오 정보 추출
        info_dict = ydl.extract_info(video_id, download=False)
        # 제목 가져오기
        title = info_dict.get('title', None)
        return title

#여기서 부터는 영어 한국어 말고 다른거도 가능하게 바꿔주려고 Ko 도 바꿔야함
@st.cache_data
def get_best_want_no_time(video_id ,learn_code,_transcript_list):
        # 선호하는 영어 자막 코드 목록
        
        try:
           # transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도)
            try:      
                transcript = _transcript_list.find_manually_created_transcript(learn_code)
                print(f"수동 생성된 배울 자막을 찾았습니다. (언어 코드: {learn_code})")
                result = [(entry['text']) for entry in transcript.fetch()]
                return result
            
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript = _transcript_list.find_generated_transcript(learn_code)
                    print(f"자동 생성된 배울 자막을 찾았습니다. (언어 코드: {learn_code})")
                    result = [(entry['text']) for entry in transcript.fetch()]
                    return result
                    
                except:
                    # 3. 마지막으로 'en'으로 시작하는 모든 코드 확인
                    available_transcripts = list(_transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith(learn_code):
                            print(f"영어 자막을 찾았습니다. (언어 코드: {learn_code})")
                        result = [(entry['text']) for entry in transcript.fetch()]
                        return result
                    
                    print("영어 자막을 찾을 수 없습니다.")
                    return None
                    
        except Exception as e:
            print(f"자막을 가져오는 중 오류가 발생했습니다: {str(e)}")
            return None

# 구글 자막을 이용해서 제미니 API가 단시간 반복 사용될때 생기는 오류를 보완하고자 하였다 하지만 해결해서 사용안함        
@st.cache_data
def get_best_to_translate_target(video_id,native_code,_transcript_list):
            
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도)
            try:
                transcript =  _transcript_list.find_manually_created_transcript(native_code)
                print(f"수동 생성된 영어 자막을 찾았습니다. {native_code})")
                result = [(entry['text']) for entry in transcript.translate(native_code).fetch()]
                
                return result
            
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript =  _transcript_list.find_generated_transcript(native_code)
                    print(f"자동 생성된 영어 자막을 찾았습니다.{native_code}")
                    result = [(entry['text']) for entry in transcript.translate(native_code).fetch()]
                    return result
                    
                except:
                    
                    available_transcripts = list(_transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith(native_code):
                            print(f" 자막을 찾았습니다.  {native_code})")
                        result = [(entry['text']) for entry in transcript.translate(native_code).fetch()]
                        return result
                    
                    print("한글 자막을 찾을 수 없습니다.")
                    return ""
 
@st.cache_data
def get_best_want_in_time(video_id ,learn_code, _transcript_list):
        # 선호하는 영어 자막 코드 목록
        
        try:
            #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도)
            try:      
                transcript =  _transcript_list.find_manually_created_transcript(learn_code)
                print(f"수동 생성된 배울 자막을 찾았습니다. {learn_code} )")
                result = [(entry['text']) for entry in transcript.fetch()]
                return result
            
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript =  _transcript_list.find_generated_transcript(learn_code)
                    print(f"자동 생성된 배울 자막을 찾았습니다. {learn_code} )")
                    result = [(entry['start'] ,entry['text']) for entry in transcript.fetch()]
                    return result
                    
                except:
                    # 3. 마지막으로 'en'으로 시작하는 모든 코드 확인
                    available_transcripts = list( _transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith('en'):
                            print(f" 배울 자막을 찾았습니다. {learn_code} )")
                        result = [(entry['start'] ,entry['text']) for entry in transcript.fetch()]
                        return result
                    
                    print("영어 자막을 찾을 수 없습니다.")
                    return None
                    
        except Exception as e:
            print(f"자막을 가져오는 중 오류가 발생했습니다: {str(e)}")
            return None
@st.cache_data
def get_best_learn_code(video_id , learn_code, _transcript_list):
        

        # 선호하는 영어 자막 코드 목록
            #transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            
            # 1. 먼저 수동 생성된 영어 자막 찾기 (모든 영어 변형 시도)
            try:
                transcript =  _transcript_list.find_manually_created_transcript(learn_code)
                print(f"수동 생성된 영어 자막을 찾았습니다.{learn_code} ")
                return learn_code
                
            except:
                # 2. 수동 자막이 없을 경우, 자동 생성된 영어 자막 찾기
                try:
                    transcript =  _transcript_list.find_generated_transcript(learn_code)
                    print(f"자동 생성된 영어 자막을 찾았습니다. {learn_code} ")
                    return learn_code
                    
                except:
                    # 3. 마지막으로 'en'으로 시작하는 모든 코드 확인
                    available_transcripts = list( _transcript_list)
                    for transcript in available_transcripts:
                        if transcript.language_code.startswith(learn_code):
                            print(f"자막을 찾았습니다. (언어 코드: {learn_code} ")
                            return learn_code
                    
                    print(" 자막을 찾을 수 없습니다.")
                    return None








def contains_lowercase(eng_script_no_time):
        """소문자 포함 여부 확인"""      
        return any(c.islower() for c in eng_script_no_time)
def check_dot(eng_script):
    """마침표 포함 여부 확인 (마침표가 5개 이하일 경우 False 반환)"""
    # 입력된 텍스트에서 마침표의 개수를 확인
    dot_count = sum(text.count('.') for text in eng_script)
    ja_dot_count = sum(text.count('。') for text in eng_script)
    # 마침표가 5개 이하라면 False, 그렇지 않으면 True 반환
    return dot_count > 5 or ja_dot_count > 5

@st.cache_data
def gemini_check_advanced_word_im_japan(_model, result_want_transcript, generation_config):
    """
    テキストを一定のサイズに分割し、各部分に対して高度な単語の説明を日本語で生成します。

    パラメータ:
        model: デルオブジェクト (API呼び出し用)
        result_eng_transcript: 全文 (文字列)
        generation_config: 生成設定 (genai.types.GenerationConfig オブジェクト)
        chunk_size: 各チャンクのサイズ (デフォルト: 300)

    戻り値:
        adw: 高度な単語の説明が含まれるリスト (各項目は文字列)
    """
    # 텍스트를 chunk_size 크기로 분할
    chunk_size = 120
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # 모델을 호출하여 고급 단어 설명 생성
        try:
            response = model.generate_content(
                f"""{chunk} テキスト内に高度な単語が含まれている場合、次の形式で単語の説明を提供してください：

                **[文中の時間] 単語 : 意味** - 文脈説明
                - '時間'は文中に単語が登場する時刻を示し、'意味'には単語の意味が入ります。
                - '文脈説明'はその単語が使用されている文脈での意味を要約します。
                例示:
                **[02:12] confirmed : 確認する - トランプがクリスティ・ノエムを国土安全保障省長官に任命するという事実を確認したという意味で使用されました。
                この形式に従ってテキスト内の高度な単語の説明を作成してください。""",
                generation_config=generation_config
            )
            # 응답 텍스트를 response_text 리스트에 추가
            response_text.append(response.text)
            response_text.append("\n")

        except Exception as e:
            print(f"오류가 발생하여 이 문장은 건너뜁니다: {e}")
            # 오류가 발생한 문장을 무시하고 다음으로 진행

    # 고급 단어 설명을 한 문자열로 합친 후, 각 줄로 분리하여 리스트로 반환
    advanced_word = "".join(response_text)
    
    return advanced_word
@st.cache_data
def gemini_check_advanced_word_im_china(_model, result_want_transcript, generation_config):
    """
    将文本按指定大小进行分割，并为每部分生成中国语的高级词汇解释。

    参数:
        model: 模型对象 (用于 API 调用)
        result_eng_transcript: 全文 (字符串)
        generation_config: 生成配置 (genai.types.GenerationConfig 对象)
        chunk_size: 每个分块的大小 (默认: 300)

    返回值:
        adw: 包含高级词汇解释的列表 (每个项目为字符串)
    """
    # 将文本分割为 chunk_size 大小的块
    chunk_size = 120
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # 调用模型生成高级词汇解释
        try:
            response = model.generate_content(
                f"""{chunk} 如果文本中包含高级词汇，请按以下格式提供词汇解释：

                **[文本中的时间] 词汇 : 含义** - 上下文解释
                - '时间'表示词汇出现在文本中的时间，'含义'提供词汇的含义。
                - '上下文解释'总结了该词汇在当前上下文中的含义。
                例句:
                **[02:12] confirmed : 确认 - 这意味着确认了特朗普将任命克里斯蒂·诺埃姆为国土安全部部长的事实。
                
                请按照此格式为文本中的高级词汇生成解释。""",
                generation_config=generation_config
            )
            # 将响应文本添加到 response_text 列表中
            response_text.append(response.text)
            response_text.append("\n")

        except Exception as e:
            print(f"出现错误，此句子将被跳过: {e}")
            # 忽略出现错误的句子并继续处理下一个句子

    # 将高级词汇解释合并为一个字符串，并按行分割为列表返回
    advanced_word = "".join(response_text)
    
    return advanced_word
@st.cache_data     
def gemini_check_advanced_word_im_fran(_model, result_want_transcript, generation_config):
    """
    Divise le texte en parties de taille spécifiée et génère des explications en français pour les mots avancés dans chaque partie.

    Paramètres:
        model: Objet du modèle (pour l'appel API)
        result_eng_transcript: Texte complet (chaîne)
        generation_config: Configuration de génération (objet genai.types.GenerationConfig)
        chunk_size: Taille de chaque morceau (par défaut : 300)

    Retour:
        adw: Liste contenant les explications des mots avancés (chaque élément est une chaîne de caractères)
    """
    # Diviser le texte en morceaux de taille chunk_size
    chunk_size = 120
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # Appeler le modèle pour générer les explications des mots avancés
        try:
            response = model.generate_content(
                f"""{chunk} Si le texte contient des mots avancés, veuillez fournir une explication pour chaque mot selon le format suivant :

                **[Temps dans le texte] Mot : Signification** - Explication contextuelle
                - 'Temps' indique le moment où le mot apparaît dans le texte, 'Signification' donne le sens du mot.
                - 'Explication contextuelle' résume le sens du mot dans le contexte où il est utilisé.
                Exemple :
                **[02:12] confirmed : Confirmer - Cela signifie avoir confirmé le fait que Trump allait nommer Kristi Noem au poste de secrétaire à la Sécurité intérieure.
                Veuillez générer des explications pour les mots avancés dans le texte en suivant ce format.""",
                generation_config=generation_config
            )
            # Ajouter le texte de la réponse à la liste response_text
            response_text.append(response.text)
            response_text.append("\n")

        except Exception as e:
            print(f"Erreur rencontrée, cette phrase sera ignorée : {e}")
            # Ignorer la phrase avec l'erreur et passer à la suivante

    # Combiner les explications des mots avancés en une seule chaîne, puis diviser chaque ligne en liste pour le retour
    advanced_word = "".join(response_text)
    
    return advanced_word
@st.cache_data
def gemini_check_advanced_word_im_espanol(_model, result_want_transcript, generation_config):
    """
    Divide el texto en partes de tamaño especificado y genera explicaciones en español para las palabras avanzadas en cada parte.

    Parámetros:
        model: Objeto del modelo (para la llamada a la API)
        result_eng_transcript: Texto completo (cadena)
        generation_config: Configuración de generación (objeto genai.types.GenerationConfig)
        chunk_size: Tamaño de cada fragmento (por defecto: 300)

    Retorno:
        adw: Lista que contiene explicaciones de palabras avanzadas (cada elemento es una cadena de caracteres)
    """
    # Dividir el texto en fragmentos de tamaño chunk_size
    chunk_size = 120
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # Llamar al modelo para generar explicaciones de palabras avanzadas
        try:
            response = model.generate_content(
                f"""{chunk} Si el texto contiene palabras avanzadas, proporcione una explicación para cada palabra según el siguiente formato:

                **[Tiempo en el texto] Palabra : Significado** - Explicación contextual
                - 'Tiempo' indica el momento en el que aparece la palabra en el texto, 'Significado' proporciona el significado de la palabra.
                - 'Explicación contextual' resume el significado de la palabra en el contexto en el que se utiliza.
                Ejemplo:
                **[02:12] confirmed : Confirmar - Esto significa haber confirmado que Trump nombrará a Kristi Noem como secretaria de Seguridad Nacional.
                Genere explicaciones para las palabras avanzadas en el texto siguiendo este formato.""",
                generation_config=generation_config
            )
            # Agregar el texto de la respuesta a la lista response_text
            response_text.append(response.text)
            response_text.append("\n")

        except Exception as e:
            print(f"Se encontró un error, se omitirá esta frase: {e}")
            # Ignorar la frase con el error y pasar a la siguiente

    # Combinar las explicaciones de palabras avanzadas en una sola cadena, luego dividir cada línea en lista para el retorno
    advanced_word = "".join(response_text)
    
    return advanced_word
@st.cache_data
def gemini_check_advanced_word(_model,result_want_transcript, generation_config):
    """
    텍스트를 일정 크기로 분할하고 각 조각에 대해 고급 단어 설명을 생성합니다.
    
    Parameters:
        model: 모델 객체 (API 호출용)
        result_eng_transcript: 전체 텍스트 (string)
        generation_config: 생성 설정 (genai.types.GenerationConfig 객체)
        chunk_size: 각 청크의 크기 (기본값: 300)
    
    Returns:
        adw: 고급 단어 설명이 포함된 리스트 (각 항목은 문자열)
    """
    # 텍스트를 chunk_size 크기로 분할
    chunk_size = 120 
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # 모델을 호출하여 고급 단어 설명 생성
        try:
            response = model.generate_content(
                f"""{chunk} 텍스트 내에 고급 단어가 포함된 경우, 다음 형식으로 단어 설명을 제공해 주세요:

                **[문장 내 시간] 단어 : 뜻** - 문맥 설명
                - '시간'은 문장에서 단어가 포함된 시간을 나타내며, '뜻'에는 단어의 의미가 들어갑니다.
                - '문맥 설명'은 해당 단어가 사용된 맥락에서 어떤 의미를 전달하는지를 요약합니다.

                예시:
                **[02:12] confirmed : 확인하다** - 트럼프가 크리스티 노임을 국토안보부 장관으로 임명할 것이라는 사실을 확인했다는 의미로 사용되었습니다

                이 형식을 참고하여 텍스트 내 고급 단어들의 설명을 작성해 주세요.""",
                generation_config=generation_config
            )
            # 응답 텍스트를 response_text 리스트에 추가
            response_text.append(response.text)
            response_text.append("\n")
        
        except Exception as e:
            print(f"오류가 발생하여 이 문장은 건너뜁니다: {e}")
            # 오류가 발생한 문장을 무시하고 다음으로 진행
            

    # 고급 단어 설명을 한 문자열로 합친 후, 각 줄로 분리하여 리스트로 반환
    advanced_word = "".join(response_text)
    
    return advanced_word 
@st.cache_data
def gemini_translate_text(_model, result_want_transcript, generation_config):
                                    
    """
    텍스트를 일정 크기로 분할하고 각 조각에 번역을 생성합니다.
    
    Parameters:
        model: 모델 객체 (API 호출용)
        result_eng_transcript: 전체 텍스트 (string)
        generation_config: 생성 설정 (genai.types.GenerationConfig 객체)
        chunk_size: 각 청크의 크기 (기본값: 300)
    
    Returns:
        translated_text: 번역된 텍스트가 포함된 문자열
    """
    # 텍스트를 chunk_size 크기로 분할
    chunk_size= 200
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # 모델을 호출하여 번역 생성
        try:
            response = model.generate_content(
                f"""
                        당신은 전문 번역가입니다. 다음 텍스트를 번역해주세요.
                        
                       

                        원본 텍스트: " {chunk}"

                        번역 시 다음 사항을 준수해주세요:
                        1. 원문의 의미를 정확하게 전달하되, 자연스러운 표현을 사용하세요.
                        2. 전문 용어가 있다면 해당 분야에서 통용되는 정확한 용어를 사용하세요.
                        3. 문화적 맥락을 고려하여 적절한 표현으로 변환하세요.
                        4. 존댓말이나 격식체의 수준을 원문과 동일하게 유지하세요.
                        
                        다음 형식으로 작성해주세요
                        
                        [문장 내 시간] 번역
                        예시
                        [00:03] 일찍이 저는 쿠르트  통 전 미국 APEC 대사이자 아시아 그룹의 매니징 파트너와 이야기를 나눴습니다.
    

                        추가 지침:
                        - 번역문만 출력하세요.
                        - 설명이나 주석을 추가하지 마세요.
                        
                        """,
                generation_config=generation_config
            )
            # 응답 텍스트를 response_text 리스트에 추가
            response_text.append(response.text)
            response_text.append("\n")
        
        except Exception as e:
            print(f"오류가 발생하여 이 문장은 건너뜁니다: {e}")
            # 오류가 발생한 문장을 무시하고 다음으로 진행

    # 번역된 텍스트를 한 문자열로 합쳐서 반환
    translated_text = "".join(response_text)
    
    return translated_text 
@st.cache_data
def gemini_translate_text_im_japan(_model, result_want_transcript, generation_config):
    """
    テキストを一定のサイズに分割し、各部分について翻訳を生成します。
    
    パラメータ:
        model: モデルオブジェクト (API呼び出し用)
        result_eng_transcript: 全体のテキスト (文字列)
        generation_config: 生成設定 (genai.types.GenerationConfig オブジェクト)
        chunk_size: 各チャンクのサイズ (デフォルト: 300)
    
    戻り値:
        translated_text: 翻訳されたテキストを含む文字列
    """
    # テキストをchunk_sizeの大きさで分割
    chunk_size = 200
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # モデルを呼び出して翻訳を生成
        try:
            response = model.generate_content(
                f"""
                あなたは専門の翻訳者です。次のテキストを翻訳してください。

                原文: "{chunk}"

                翻訳の際、以下の点に従ってください:
                1. 原文の意味を正確に伝えつつ、自然な表現を使用してください。
                2. 専門用語が含まれる場合、その分野で一般的に使用される正確な用語を使用してください。
                3. 文化的な文脈を考慮して、適切な表現に変換してください。
                4. 敬語や丁寧語のレベルを原文と同じにしてください。

                「次の形式で書いてください」

                [文中の時間] 翻訳
                例え
                [00:03] 以前、私は元アメリカAPEC大使であり、アジアグループのマネージングパートナーであるクルト・トング氏とお話ししました。

                追加指示:
                - 翻訳文のみ出力してください。
                - 説明や注釈を追加しないでください。
                """,
                generation_config=generation_config
            )
            # 応答テキストをresponse_textリストに追加
            response_text.append(response.text)
            response_text.append("\n")
        
        except Exception as e:
            print(f"エラーが発生したため、この文はスキップされます: {e}")
            # エラーが発生した文を無視して次に進む

    # 翻訳されたテキストを1つの文字列に結合して返す
    translated_text = "".join(response_text)
    
    return translated_text
@st.cache_data
def gemini_translate_text_im_china(_model,result_want_transcript, generation_config):
    """
    テキストを一定のサイズに分割し、各部分について翻訳を生成します。
    
    パラメータ:
        model: モデルオブジェクト (API呼び出し用)
        result_eng_transcript: 全体のテキスト (文字列)
        generation_config: 生成設定 (genai.types.GenerationConfig オブジェクト)
        chunk_size: 各チャンクのサイズ (デフォルト: 300)
    
    戻り値:
        translated_text: 翻訳されたテキストを含む文字列
    """
    # テキストをchunk_sizeの大きさで分割
    chunk_size = 200
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # モデルを呼び出して翻訳を生成
        try:
            response = model.generate_content(
                f"""
                あなたは専門の翻訳者です。次のテキストを翻訳してください。

                原文: "{chunk}"

                翻訳の際、以下の点に従ってください:
                1. 原文の意味を正確に伝えつつ、自然な表現を使用してください。
                2. 専門用語が含まれる場合、その分野で一般的に使用される正確な用語を使用してください。
                3. 文化的な文脈を考慮して、適切な表現に変換してください。
                4. 敬語や丁寧語のレベルを原文と同じにしてください。
                
                「请按照以下格式书写」
                [句子中的时间] 翻译

                例子
                [00:03] 此前，我与前美国APEC大使、亚洲集团的管理合伙人库尔特·通先生进行了交谈。
                
                追加指示:
                - 翻訳文のみ出力してください。
                - 説明や注釈を追加しないでください。
                """,
                generation_config=generation_config
            )
            # 応答テキストをresponse_textリストに追加
            response_text.append(response.text)
            response_text.append("\n")
        
        except Exception as e:
            print(f"エラーが発生したため、この文はスキップされます: {e}")
            # エラーが発生した文を無視して次に進む

    # 翻訳されたテキストを1つの文字列に結合して返す
    translated_text = "".join(response_text)
    
    return translated_text
@st.cache_data
def gemini_translate_text_im_espanol(_model, result_want_transcript, generation_config):

    """
    Divide el texto en partes de un tamaño específico y genera la traducción para cada fragmento.
    
    Parámetros:
        model: objeto del modelo (para la llamada API)
        result_eng_transcript: texto completo (cadena de texto)
        generation_config: configuración de generación (objeto genai.types.GenerationConfig)
        chunk_size: tamaño de cada fragmento (predeterminado: 300)
    
    Retorna:
        translated_text: cadena de texto que contiene el texto traducido
    """
    # Divide el texto en fragmentos de tamaño chunk_size
    chunk_size = 200
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # Llama al modelo para generar la traducción
        try:
            response = model.generate_content(
                f"""
                Eres un traductor profesional. Traduce el siguiente texto.

                Texto original: "{chunk}"

                Al traducir, sigue estos puntos:
                1. Transmite el significado del texto original con precisión y utiliza una expresión natural.
                2. Si contiene términos técnicos, utiliza el término correcto que se usa comúnmente en el campo correspondiente.
                3. Considera el contexto cultural y conviértelo a una expresión adecuada.
                4. Mantén el nivel de cortesía o formalidad igual al del texto original.

                Por favor, escríbalo en el siguiente formato.

                [Tiempo en la oración] Traducción

                ejemplo
                [00:03] Anteriormente, hablé con Kurt Tong, ex embajador de EE. UU. en APEC y socio gerente del Grupo Asia.
                
                Instrucciones adicionales:
                - Solo muestra el texto traducido.
                - No añadas explicaciones ni notas.
                """,
                generation_config=generation_config
            )
            # Agrega el texto de respuesta a la lista response_text
            response_text.append(response.text)
            response_text.append("\n")
        
        except Exception as e:
            print(f"Se produjo un error, esta oración se omitirá: {e}")
            # Ignora la oración en la que se produjo un error y continúa

    # Une el texto traducido en una sola cadena y lo retorna
    translated_text = "".join(response_text)
    
    return translated_text
@st.cache_data
def gemini_translate_text_im_fran(_model,result_want_transcript, generation_config): 
    """
    Divise le texte en parties de taille spécifique et génère la traduction pour chaque fragment.
    
    Paramètres :
        model : objet du modèle (pour l'appel API)
        result_eng_transcript : texte complet (chaîne de caractères)
        generation_config : configuration de génération (objet genai.types.GenerationConfig)
        chunk_size : taille de chaque fragment (par défaut : 300)
    
    Retourne :
        translated_text : chaîne de caractères contenant le texte traduit
    """
    # Divise le texte en fragments de taille chunk_size
    chunk_size = 200
    chunks_script = [result_want_transcript[i:i + chunk_size] for i in range(0, len(result_want_transcript), chunk_size)]
    response_text = []

    for chunk in chunks_script:
        # Appelle le modèle pour générer la traduction
        try:
            response = model.generate_content(
                f"""
                Vous êtes un traducteur professionnel. Traduisez le texte suivant.

                Texte original : "{chunk}"

                Lors de la traduction, veuillez respecter les points suivants :
                1. Transmettez précisément le sens du texte original tout en utilisant une expression naturelle.
                2. Si des termes techniques sont inclus, utilisez le terme correct couramment employé dans le domaine concerné.
                3. Prenez en compte le contexte culturel et adaptez l'expression de manière appropriée.
                4. Maintenez le même niveau de politesse ou de formalité que dans le texte original.

                Veuillez l'écrire dans le format suivant
                [Temps dans la phrase] Traduction

                exemple
                [00:03] Auparavant, j'ai parlé avec Kurt Tong, ancien ambassadeur des États-Unis auprès de l'APEC et partenaire directeur du Groupe Asie.
                
                Instructions supplémentaires :
                - Affichez uniquement le texte traduit.
                - N'ajoutez ni explications ni notes.
                """,
                generation_config=generation_config
            )
            # Ajoute le texte de réponse à la liste response_text
            response_text.append(response.text)
            response_text.append("\n")
        
        except Exception as e:
            print(f"Une erreur est survenue, cette phrase sera ignorée : {e}")
            # Ignore la phrase ayant causé une erreur et continue

    # Concatène le texte traduit en une seule chaîne et la retourne
    translated_text = "".join(response_text)
    
    return translated_text


import streamlit as st


def create_settings_sidebar():
    with st.sidebar:
        st.title("🛠️ SETTING")
        #st.subheader("Developer's blog")
        
        # st.markdown("""
        #     <a href="https://fktshin.tistory.com/15" 
        #        target="_blank"
        #        style="text-decoration: none;">
        #         <div style="background-color: #A8E6CF; 
        #                     padding: 7px 20px; 
        #                     border-radius: 10px; 
        #                     text-align: center; 
        #                     margin: 10px 0;
        #                     display: inline-block;
        #                     width: 100%;
        #                     box-sizing: border-box;
        #                     font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        #                     box-shadow: 0 1px 3px rgba(0,0,0,0.12), 0 1px 2px rgba(0,0,0,0.24);">
        #             <div style="display: flex; 
        #                         align-items: center; 
        #                         justify-content: center; 
        #                         gap: 8px;
        #                         color: #1a1a1a;
        #                         font-size: 14px;
        #                         font-weight: 500;
        #                         letter-spacing: 0.25px;">
        #                 <span style="color: #1a1a1a;">🌱</span>
        #                 <span style="color: #1a1a1a;">Visit Blog!</span>
        #             </div>
        #         </div>
        #     </a>
        #     """, unsafe_allow_html=True)
        # Gemini API 설정 섹션
        st.header("API Setting")
        api_key = st.text_input(
            "Google Gemini 1.5 Flash API Key",
            type="password",
            help="Please enter the Gemini API key issued from Google Cloud Console."
        )
         # API 키 발급 링크 버튼
        st.markdown("""
            <a href="https://aistudio.google.com/apikey?hl=ko&_gl=1*1841s3x*_ga*MTk1Nzc2OTYwMi4xNzMwOTQyNDU0*_ga_P1DBVKWT6V*MTczMTcyOTY1MC4xMC4wLjE3MzE3Mjk2NTAuNjAuMC4xODI4NzE1NDg." 
               target="_blank"
               style="text-decoration: none;">
                <div style="background-color: #619BF7; 
                            padding: 10px 20px; 
                            border-radius: 25px; 
                            text-align: center; 
                            margin: 10px 0;
                            display: inline-block;
                             width: 100%;
                            box-sizing: border-box;
                            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                            box-shadow: 0 1px 3px rgba(0,0,0,0.12), 0 1px 2px rgba(0,0,0,0.24);">
                    <div style="display: flex; 
                                align-items: center; 
                                justify-content: center; 
                                gap: 8px;
                                color: #1a1a1a;
                                font-size: 14px;
                                font-weight: 500;
                                letter-spacing: 0.25px;">
                        <span style="color: #1a1a1a;">🔑</span>
                        <span style="color: #1a1a1a;">Get your Free Api key.</span>
                    </div>
                </div>
            </a>
            """, unsafe_allow_html=True)
        
        st.header("Font size Setting")
        
        # 모국어와 학습 언어의 폰트 설정
        native_font = st.text_input(
            "Native Language font",
            value="Noto Sans",
            help="default setting : Noto Sans"
        )
        
        want_font = st.text_input(
            "Target Language Font",
            value="Cambria",
            help="default setting : Cambria"
        )
        
        # 폰트 크기 설정
        font_size = st.slider(
            "Font Size (px)",
            min_value=8,
            max_value=20,
            value=11,
            step=1
        )

        # st.subheader("Developer's blog")
        # st.markdown("""
        #     <a href="https://fktshin.tistory.com/15" 
        #        target="_blank"
        #        style="text-decoration: none;">
        #         <div style="background-color: #A8E6CF; 
        #                     padding: 7px 20px; 
        #                     border-radius: 10px; 
        #                     text-align: center; 
        #                     margin: 10px 0;
        #                     display: inline-block;
        #                     width: 100%;
        #                     box-sizing: border-box;
        #                     font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        #                     box-shadow: 0 1px 3px rgba(0,0,0,0.12), 0 1px 2px rgba(0,0,0,0.24);">
        #             <div style="display: flex; 
        #                         align-items: center; 
        #                         justify-content: center; 
        #                         gap: 8px;
        #                         color: #1a1a1a;
        #                         font-size: 14px;
        #                         font-weight: 500;
        #                         letter-spacing: 0.25px;">
        #                 <span style="color: #1a1a1a;">🌱</span>
        #                 <span style="color: #1a1a1a;">Visit Blog!</span>
        #             </div>
        #         </div>
        #     </a>
        #     """, unsafe_allow_html=True)
        
        return {
            "api_key": api_key,
            "native_font": native_font,
            "want_font": want_font,
            "font_size": font_size,
        }

settings = create_settings_sidebar()

if not settings["api_key"]:
        st.warning("Please enter the Gemini API key in the sidebar.")
        st.stop()
# st.write(f"- Slected native language font: {settings['native_font']}")
# st.write(f"- Slected target language font: {settings['want_font']}")
# st.write(f"- Font Size: {settings['font_size']}px")

native_font = settings['native_font']
want_font = settings['want_font']
font_size = settings['font_size']
api_key = settings['api_key']

user_input= url
video_id = get_video_id(user_input) 

if video_id is None :
    display_chat_message("assistant", "Please check the URL address again.")
    
else:
    try:
        # transcript_list 초기화
       
        title_video = get_video_title(video_id)
        try:
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        except:
           st.warning("Google banned this service...{e}")
        
        print("api_work")
                                   
        # 채팅 처리
        if user_input:
            # 사용자 메시지 표시
            # st.markdown("""<a href="https://link.coupang.com/a/b4Z0hb" target="_blank" referrerpolicy="unsafe-url">
            #             <img src="https://img4a.coupangcdn.com/image/affiliate/widget/image/2024/11/25/03a1609c92ee009c010ecaed3566ed44.png" 
            #             alt=""style=" margin-bottom:15px"></a>""",unsafe_allow_html=True)
           
            #display_chat_message("user", user_input)

            ai_response = "Checking your video... Bot is running!"

            try:
                # 자막 언어 목록 가져오기
                    display_chat_message("assistant", ai_response)
    
                #if want_language == "English":# 영어가 배우고 싶은 사람이 많을테니 힘을 썼다.
                    

                    if want_language == "English":

                        want_lang_no_time = get_best_english_transcript_no_time(video_id,transcript_list)
                        want_code_check = get_best_english_encode(video_id,transcript_list)
                        want_lang_in_time = get_best_english_transcript(video_id,transcript_list)
                       
                        
                    else:
                        want_lang_no_time = get_best_want_no_time(video_id, learn_code,transcript_list)
                        want_code_check  = get_best_learn_code(video_id , learn_code,transcript_list)
                        want_lang_in_time= get_best_want_in_time(video_id,learn_code,transcript_list)
                       
                    
                    if want_code_check == None:
                        display_chat_message("assistant" , "There are no subtitles available. Please choose a different video.")
                        

                    else:                                     
                           
                            UporLow = contains_lowercase(want_lang_no_time)
                            dot_Check = check_dot(want_lang_no_time)
                        
                            if dot_Check == False:
                                if want_language  != "English":                                      
                                    #display_chat_message("assistant", "There are no delimiters. This service analyzes sentences using delimiters such as (.) and (?). Please select a different video.")
                                    st.warning("There are no delimiters. This service analyzes sentences using delimiters such as (.) and (?). Please select a different video.")
                                    st.stop()
                    
                            
                            import re 

                            new_script = ""

                            for start, read_script in want_lang_in_time:
                                    minutes = int(start // 60)  # 분 계산 (소수점 없음)
                                    seconds = int(start % 60)  # 초 계산 (소수점 없음)

                                    if dot_Check == False and want_language  == "English": # 영어 일때만 문장 구분 
                                        #display_chat_message("assistant", dot_Check)
                                            # 문장 구분이 필요한 단어 리스트
                                        keyword = [
                                              # 전치사 및 접속사 (대문자로 시작)
                                        "And ", "But ", "What", "How", "Who", "Have", "Did", "If ", "When", "Because", 
                                        "Then", "Or", "Why", "Although", "Though", "After", "Before", "Until", 
                                        "While", "Since", "Once", "Thus", "Hence", "Consequently", "Unless", "Whether", 
                                        "Therefore", "However", "Moreover", "Yet", "Still", "Nonetheless", "Nevertheless", 
                                        "For instance", "Such as", "For example", "On the other hand", "To summarize", "Otherwise",
                                    
                                    ]
                                        for word in keyword:
                                            read_script  = read_script.replace(word , f".{word}")
                                        
                                        
                                     
                                
                                # 시간 형식 설정 (분.초 형태)
                                    time_format = f"[{minutes:02d}:{seconds:02d}]"
                                # . 기반이다 보니 문제가 있을 만한 것들을 수정    
                                    replacements = {
                                        'U.S.': 'US',
                                        'U.S': 'US',
                                        'S.E.C.': 'SEC',
                                        'Mr.': 'Mr ',
                                        'Mrs.': 'Mrs ',
                                        'Ph.D.': 'PhD ',
                                        'Prof.': 'prof ',
                                        'Dr.': 'Dr ',
                                        'No.': 'Number ',
                                        'a.m.': 'am',
                                        'p.m.': 'pm',
                                        '.com': '(dot)com',
                                        '.site': '(dot)site'
                                    }

                                    # 치환 반복 적용
                                    for old, new in replacements.items():
                                        read_script = read_script.replace(old, new)
                                    
                                    read_script = re.sub(r'(\d)\.(\d)', r'\1 point \2', read_script)


                                    read_script = read_script.replace('..' , ".")
                                    read_script = read_script.replace('..' , ".")
                                    read_script = read_script.replace('..' , ".") 
                                    

                                    read_script = read_script.replace('\n', ' ')
                                    read_script = read_script.replace('.', '. \n')
                                    read_script = read_script.replace('?' , '? \n')
                                    read_script = read_script.replace('。' , '。 \n')
                                   
                                    if UporLow  == False:
                                            read_script = read_script[0].upper() + read_script[1:].lower()

                                            
 
                                    new_script += ' '
                                    new_script += time_format
                                    new_script += read_script

                                      
                            result_want_transcript =  ["\n\n"]

                          

                            # 타임스탬프 맨앞 빼고 제거 함수
                            def clean_transcript_texts(transcript_texts):
                                        cleaned_texts = ""
                                        for text in transcript_texts:
                                            # 첫 번째 타임스탬프만 남기고 나머지 타임스탬프를 제거
                                            # 1) 모든 타임스탬프를 찾음
                                            timestamps = re.findall(r'\[\d{2}:\d{2}\]', text)
                                            
                                            if timestamps:      
                                                first_timestamp = timestamps[0]

                                                cleaned_text = text.replace(first_timestamp, '',1)
                                                cleaned_text = re.sub(r'\[\d{2}:\d{2}\]','', cleaned_text)  # 나머지 타임스탬프 제거
                                                cleaned_texts += first_timestamp +" "+ cleaned_text.strip() +" "
                                            
                                            else:
                                                # 타임스탬프가 없는 경우
                                                cleaned_texts += text.strip() +" "

                                        return cleaned_texts.strip()  # 최종 문자열 반환

                            result_want_script = new_script.splitlines()
                            
                            
                            result_if_too_Long = ""


                            keywords = [
                            # 전치사 및 접속사
                            # "And ", "and ", "But ", "but ", "What", "what", "How", "how", "Who","who",
                            # "Have", "Did", "did", "If ", "if ", "When", "when", "Because", "because",
                            # "Then", "then", "Or", "Why", "why", "Although", "although", 
                            # "Though", "though", "After", "after", "Before", "Until", 
                            # "While", "while", "Since", "since", "Once", "once", "Thus", "thus", "Hence", "hence", 
                            # "Consequently", "Unless", "unless", "Whether", "whether", "Therefore", 
                            # "therefore", "However", "however", "Moreover", "moreover", "Yet", "yet", "Still", 
                            # "Nonetheless", "Nevertheless" , "For instance", "for instance", 
                            # "Such as", "For example", "for example", "On the other hand", 
                            # "on the other hand", "To summarize", "to summarize","Otherwise","otherwise"
                            # 전치사 및 접속사 (소문자로 시작)
                            "and ", "but ", "what", "how ", "who ", "did ", "if ", "when", "because", 
                            "then", "or", "why", "although", "though", "after", "before", "while", 
                            "since", "once", "thus", "hence", "unless", "whether", "therefore", 
                            "however", "moreover", "yet", "still", "for instance", "such as", 
                            "for example", "on the other hand", "to summarize", "otherwise","or "
                            #대명사 및 기타 단어 (대문자로 시작)
                            "I ","I'm" "I've", "We", "You", "They", "It ","It's", "He ", "She ", 
                            "This ", "That ", "These", "Those", 
                            "Indeed", "Certainly", "Surely", "Alright", "There ",

                            # 대명사 및 기타 단어 (소문자로 시작)
                            "i ", "i've","i'm", "we", "you","they", "it ","it's", "he ","he's", "she ","she's" 
                            "this ", "these", "those", 
                            "indeed", "certainly", "surely", "alright", "there "
                        ]
                            pattern = r'(\b(?:' + '|'.join(map(re.escape, keywords)) + r'))'  # 모든 키워드를 포함한 정규식

                            result_if_too_Long = ""

                            for line in result_want_script:
                                if len(line) > 230:
                                    start = 0  # 검색 시작 위치
                                    
                                    while start < len(line):
                                        matches = list(re.finditer(pattern, line[start:]))  # 남은 문자열에서 모든 키워드 찾기
                                        
                                        if not matches:  # 더 이상 매칭된 키워드가 없으면 종료
                                            break
                                        
                                        found_valid = False
                                        for match in matches:
                                            keyword_start = start + match.start(1)  # 매칭된 키워드의 시작 위치
                                            keyword_end = start + match.end(1)     # 매칭된 키워드의 끝 위치

                                            # 키워드 앞에 다른 알파벳이 없는 경우만 처리
                                            if keyword_start > 0 and line[keyword_start - 1].isalpha():
                                                continue
                                            
                                            # 조건 확인: 키워드 앞뒤로 100자 이상인지
                                            if len(line[start:keyword_start]) >= 125 and len(line[keyword_end:]) >= 125:
                                                # 키워드 앞에 ".\n" 추가
                                                line = line[:keyword_start] + ".\n" + line[keyword_start:]
                                                start = keyword_end + 2  # ".\n"의 길이를 추가하여 다음 검색 시작 위치 갱신
                                                found_valid = True
                                                break  # 조건을 만족하면 다음 반복으로 이동

                                        if not found_valid:
                                            # 조건을 만족하는 키워드가 없으면 start를 다음 키워드 이후로 이동
                                            start = keyword_end

                                result_if_too_Long += line + "\n"


                            # 결과 확인
                            
                                #display_chat_message("assistant" , result_if_too_Long)
                                result_want_script = result_if_too_Long.splitlines()
                           
                            non_time_line_last = ""
                            non_time_check_last = True
                            for line in  result_want_script:
                                    
                                    time_judge = re.search(r"\[(\d{2}:\d{2})\]", line)

                                    if time_judge:
                                        
                                        if non_time_check_last == False:
                                             non_time_line_last += line
                                             line = non_time_line_last
                                             non_time_line_last=""
                                             non_time_check_last = True
                                        else:
                                            result_want_transcript.append("\n")
                                            result_want_transcript.append(clean_transcript_texts([line]))
                                            result_want_transcript.append("\n")   
                                    else:
                                        non_time_line_last = line
                                        non_time_check_last = False

                                    # result_want_transcript.append("\n")
                                    # result_want_transcript.append(clean_transcript_texts([line]))
                                    # result_want_transcript.append("\n")    

                            result_only_want_for_word = ["\n\n\n"]
                            non_time_line = ""

                            for line in result_want_script:
                                time_judge = re.search(r"\[(\d{2}:\d{2})\]", line)

                                if time_judge:
                                    if non_time_line:  # non_time_line이 비어있지 않으면
                                        line  = non_time_line + line
                                        result_only_want_for_word.append(clean_transcript_texts([non_time_line]))
                                        result_only_want_for_word.append("\n")
                                        non_time_line = ""  # non_time_line 초기화
                                    else:  
                                        result_only_want_for_word.append(clean_transcript_texts([line]))
                                        result_only_want_for_word.append("\n")
                                else:
                                    non_time_line = line  # 현재 줄을 non_time_line에 저장
                                               
                            
                            word_file = create_word_file_shadow_script(result_only_want_for_word,title_video,learn_code,want_font,native_font,font_size)
                            
    
                            #유사도 기반 아이디어를 삭제 하였음
                            #처음에는 artifact 저장소의 비용을 줄이고자 해서 삭제했지만 새로 수정한 버전이 좋은 성능을 보임

                            # result_target_script =  get_best_to_translate_target(video_id ,native_code,transcript_list) 
                            # new_target_script = ""

                            # for read_script_target_line in result_target_script:
                                    
                            #         read_script_target_line = read_script_target_line.replace('U.S.', 'US')
                            #         read_script_target_line = read_script_target_line.replace('U.S', 'US')
                            #         read_script_target_line = read_script_target_line.replace('S.E.C.' , 'SEC')
                            #         read_script_target_line = read_script_target_line.replace('Mr.', 'Mr ')
                            #         read_script_target_line = read_script_target_line.replace('Mrs.', 'Mrs ')

                            #         read_script_target_line = read_script_target_line.replace('Ph.D.', 'ph,D ')
                            #         read_script_target_line = read_script_target_line.replace('Prof.', 'prof ')
                            #         read_script_target_line = read_script_target_line.replace('Dr.', 'Dr ')

                            #         read_script_target_line = read_script_target_line.replace('No.', 'Number')
                            #         read_script_target_line = read_script_target_line.replace('a.m.', 'am')
                            #         read_script_target_line= read_script_target_line.replace('p.m.', 'pm')

                            #         read_script_target_line = read_script_target_line.replace('\n', ' ')
                            #         read_script_target_line = read_script_target_line.replace('.', '. \n')
                            #         read_script_target_line = read_script_target_line.replace('。', '。 \n')
                            #         read_script_target_line = read_script_target_line.replace('?' , '? \n')
                                    
                                    
                            #         new_target_script +=' '
                            #         new_target_script += read_script_target_line
                            
                            
                            # kor_script_line = new_target_script.splitlines()
                            
                            import google.generativeai as genai
                                                            
                            try:
                                genai.configure(api_key=api_key)
                                model = genai.GenerativeModel("gemini-1.5-flash")
                                generation_config = genai.types.GenerationConfig(
                                    candidate_count=1,
                                    stop_sequences=["x"],
                                    temperature=0,
                                )
                                response = model.generate_content("Hello, how is the API working? if you're working , put your hands UP!!")
                            except Exception as e:
                                print(f"An error occurred: {e}")
                                # Add further actions like retrying or prompting for a valid API key.
                                st.warning("Please check your Gemini API again.")
                                st.stop()

                            display_chat_message("assistant","I'm working hard on the analysis, but it might take some time. Please wait a moment!") 
                            #언어 분석은 이렇게 가자 
                            if native_code == "ja":
                                    advanced_word = gemini_check_advanced_word_im_japan(model, result_want_transcript, generation_config)
                            if native_code == "ko":
                                advanced_word = gemini_check_advanced_word(model, result_want_transcript, generation_config)

                            if native_code == "es":
                                    advanced_word = gemini_check_advanced_word_im_espanol(model, result_want_transcript, generation_config)
                            if native_code == "zh-Hans":
                                    advanced_word = gemini_check_advanced_word_im_china(model, result_want_transcript, generation_config)

                            if native_code == "fr":
                                    advanced_word =gemini_check_advanced_word_im_fran(model, result_want_transcript, generation_config)
                            
                            adw_script = advanced_word.splitlines()    
                            word_file_adw = create_word_file_shadow_script(adw_script ,title_video,learn_code,want_font,native_font,font_size)

                           
                            display_chat_message("assistant","Almost there! Just a few more moments..") 
                            #display_chat_message("assistant", advanced_word)
                            if native_code == "ja":
                                gemini_transcript= gemini_translate_text_im_japan(model, result_want_transcript, generation_config)
                            if native_code == "ko":
                                gemini_transcript= gemini_translate_text(model, result_want_transcript, generation_config)

                            if native_code == "es":
                                    gemini_transcript= gemini_translate_text_im_espanol(model, result_want_transcript, generation_config)
                            if native_code == "zh":
                                    gemini_transcript= gemini_translate_text_im_china(model, result_want_transcript, generation_config)

                            if native_code == "fr":
                                    gemini_transcript= gemini_check_advanced_word_im_fran(model, result_want_transcript, generation_config)

                           

                            #원래는 문제가 있는 개별을 번역해주려했으나 가끔 오류가 발생 사용불가
                            def translate_with_gemini(model, text, source_lang, target_lang):
                                        """
                                        Gemini를 사용하여 텍스트를 번역합니다.
                                        
                                        매개변수:
                                        model: Gemini 모델 인스턴스
                                        text (str): 번역할 텍스트
                                        source_lang (str): 원본 언어 (예: '한국어', '영어', '일본어')
                                        target_lang (str): 목표 언어 (예: '영어', '한국어', '일본어')
                                        generation_config: 모델의 생성 설정
                                        max_retries (int): 최대 재시도 횟수
                                        
                                        반환:
                                        str: 번역된 텍스트
                                        """
                                        attempt = 0
                                        while attempt < 3:
                                            try:
                                                response = model.generate_content(
                                                    f"""
                                                    당신은 전문 번역가입니다. 다음 텍스트를 {source_lang}에서 {target_lang}으로 번역해주세요.

                                                    원본 텍스트: "{text}"

                                                    번역 시 다음 사항을 준수해주세요:
                                                    1. 원문의 의미를 정확하게 전달하되, 자연스러운 표현을 사용하세요.
                                                    2. 존댓말이나 격식체의 수준을 원문과 동일하게 유지하세요.
                                                    3. 당신은 친절하고 정중하며 완벽하게 번역을 수행하는 번역가입니다. 
                                                    
                                                    번역문만 출력하세요.
                                                    """,
                                                    generation_config=generation_config
                                                )
                                                
                                                # 응답이 유효한지 확인
                                                if hasattr(response, 'text') and response.text:
                                                    return response.text
                                                else:
                                                    raise ValueError("유효하지 않은 응답입니다.")
                                            
                                            except Exception as e:
                                                print(f"오류가 발생했습니다 (시도 {attempt + 1}/3): {e}")
                                                attempt += 1

                                        # 최대 재시도 횟수 초과 시 None 반환
                                        print("최대 재시도 횟수를 초과했습니다.")
                                        return None
                    

                            from sentence_transformers import SentenceTransformer  # 텍스트 백터 변환
                            from sklearn.metrics.pairwise import cosine_similarity # 벡터 유사도 계산
                            import numpy as np

                            # 파일 읽기 리스트화 하였습니다
                            
                            english_lines = result_want_transcript

                            gemini_lines = gemini_transcript.splitlines()  # splitlines()로 리스트 생성
                           
                            #display_chat_message("assistant","If you add this timestamp to the YouTube comments, it will make studying easier.")    
                          
                            #유사도기반 없엠 클라우드 비용이 높고 , 제미니 시간 에러 적고 해결하였음     
                             #korean_lines = kor_script_line
                            # 문장 임베딩 모델 로드 (다국어 지원 모델 사용)
                            # model_simul = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')#paraphrase-xlm-r-multilingual-v1
                            # #model_simul = SentenceTransformer('paraphrase-xlm-r-multilingual-v1')
                            
                            # # 영어와 한글 문장의 임베딩 벡터 생성 # 임베딩 생성
                            # # 문장 임베딩 모델 로드 (도커 컨테이너 내부 경로에서 로드)
                            # #model_simul = SentenceTransformer('/app/model/sentence_transformer')
                            # english_embeddings = model_simul.encode(english_lines)
                            # korean_embeddings = model_simul.encode(korean_lines)

                            # 유사도 매트릭스 계산
                            #similarity_matrix = cosine_similarity(english_embeddings, korean_embeddings) 

                            # 유사도가 가장 높은 문장끼리 매칭
                            
                            


                            time_list  = []
                            merged_lines = ["\n\n\n"]
                            used_korean_indices = set() # 사용한 한국어는 지우기 위해 집합 사용
                            used_indices = set()
                            for eng_idx, eng_sentence in enumerate(english_lines):
                                # 각 영어 문장에 대해 가장 유사한 한글 문장을 찾음
                                if not eng_sentence.strip():
                                    continue

                                time_judge = re.search(r"\[(\d{2}:\d{2})\]", eng_sentence)
                                
                                if time_judge:  # time_judge가 None이 아닐 때
                                    time_str = time_judge.group(0) 
                                    
                                    time_list.append(time_str)
                                    
                                    for j,gemini_line in enumerate(gemini_lines):
                                            if j in used_indices:
                                                 continue
                                                      
                                            if time_str in gemini_line:  
                                                kor_sentence = gemini_line
                                                kor_sentence = re.sub(r'\[\d{2}:\d{2}\]','', kor_sentence)
                                                used_indices.add(j)
                                                break
                                            
                                              
                                else:
                                    # best_kor_idx = np.argmax(similarity_matrix[eng_idx])
                                    # best_kor_similarity = similarity_matrix[eng_idx, best_kor_idx]             

                                    # if best_kor_idx not in used_korean_indices:
                                                        
                                    #         kor_sentence = korean_lines[best_kor_idx]

                                    # used_korean_indices.add(best_kor_idx)
                        
                                    kor_sentence = ""
                    


                                merged_lines.append(eng_sentence)
                                merged_lines.append("\n\n")
                                merged_lines.append(kor_sentence)
                               
                                                      
                                
                                #시간이 부정확한 제미니가 있어서 그부분을 바꾸고 중복 단어도 보여주게 했다 계속보면서 공부하게        
                                #사람에 따라서 불편한점이 있을듯 제미니가 해결되면 전환해도..?

                                if time_judge: 
                                    time_str = time_judge.group(0)
                                    for j, script_line in enumerate(adw_script):
                                        if not script_line.strip():
                                            continue

                                        if time_str:
                                            adw_time_judge = re.search(r"\[(\d{2}:\d{2})\]", script_line)
                                            if adw_time_judge:
                                                adw_time_str = adw_time_judge.group(0)
                                                adw_replace_time = script_line.replace(adw_time_str, "")
                                                adw_index = adw_replace_time.find(":")
                                                                                                
                                                if adw_index != -1:
                                                   
                                                    if eng_sentence.find(adw_replace_time[2:adw_index]) != -1:
                                                      
                                                        merged_lines.append("\n")
                                                        merged_lines.append(adw_replace_time)

                                    merged_lines.append("\n\n")
                                else:
                                    merged_lines.append("\n")
                                               
                           
                                   
                                #시간 인덱스 말고 단어 인덱스로 변환하였음 제미니에 들어가고 시간 오류가 가끔 나는 경우가 있었음
                                # if time_judge:  # time_judge가 None이 아닐 때
                                #     time_str = time_judge.group(0) 
                                #     for j in range(len(adw_script)):
                                #         if time_str in adw_script[j]:  # time_str이 adw[j]에 있는지 확인
                                #             merged_lines.append("\n")
                                #             merged_lines.append(adw_script[j].replace(time_str,""))
                                                
                                #     merged_lines.append("\n\n")      
                                # else:
                                #     merged_lines.append("\n")

                            # display_chat_message("assistant","If you add this timestamp to the YouTube comments, it will make studying easier.")    
                            # display_chat_message("assistant", time_list)  
                            merged_lines.insert(0, "\n" + "✨ Adding this timestamp to YouTube comments simplifies studying! :} ✨" + "\n" + str(time_list))
                            merged_en_ko_script = "".join(merged_lines)

                                
                            merged_en_ko_script_split = merged_en_ko_script.splitlines()
                                 
                                        
                            word_file_shadowing_script = create_word_file_shadow_script(merged_en_ko_script_split,title_video,learn_code,want_font,native_font,font_size)
                                # 워드 파일 다운로드 버튼
                            
                            st.download_button(
                                            label="📄 Download Learning Languages Script.docx",
                                            data=word_file,
                                            file_name="Learning Languages Script.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key="download_button_1"
                                )
                            st.download_button(
                                            label="📄 Download Difficult Words List.docx",
                                            data=word_file_adw,
                                            file_name="Difficult_Words_List.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key="download_button_word_file_adw"
                                )
                            
                            st.download_button(
                                    label="📄 Download Shadowing_Script (Translation and Difficult Words).docx",
                                    data=word_file_shadowing_script ,
                                    file_name="Shadowing Script.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_button_merged_shadowing_script"
                                    )            
                            st.success("I've completed it! Expand your world!", icon="✅")
                            #display_chat_message("assistant","I've completed it! Expand your world!")
                            st.balloons()
                            
                            st.markdown("""
                            <iframe src="https://ads-partners.coupang.com/widgets.html?id=823342&template=carousel&trackingCode=AF4610152&subId=&width=680&height=120&tsource=" 
                                    width="680" height="100" frameborder="0" scrolling="no" referrerpolicy="unsafe-url" browsingtopics></iframe>
                            """, unsafe_allow_html=True)
                            # st.markdown("""<a href="https://link.coupang.com/a/b40cin" target="_blank" referrerpolicy="unsafe-url">
                            # <img src="https://image11.coupangcdn.com/image/cmg/oms/banner/0ebebf1e-2dc2-423b-933a-3a9f14dac987_980x150.jpg" 
                            # alt=""style=" margin-bottom:15px"></a>""",unsafe_allow_html=True)
                            
                            display_chat_message("assistant", "해당 서비스는 쿠팡 파트너스 활동을 통해 일정액의 수수료를 제공받을 수 있습니다.")
            except Exception as e:
                # list_available_languages에서 에러가 발생하면 처리
                st.warning(f"YouTube subtitles access is restricted. Please choose another video")
    except Exception as e:
        # transcript_list 초기화에서 에러가 발생하면 처리
         st.warning(f"YouTube subtitles access is restricted. Please choose another video.") 